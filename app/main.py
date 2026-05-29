from __future__ import annotations

import asyncio
from html import escape as html_escape
from io import BytesIO
import hashlib
import hmac
import logging
import json
import re
from contextlib import suppress
from pathlib import Path
from dataclasses import replace
from datetime import datetime, timedelta, timezone
from typing import Any

import uvicorn
from aiogram import Bot
from aiogram.types import InlineKeyboardButton, InlineKeyboardMarkup, WebAppInfo
from fastapi import Depends, FastAPI, Header, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse, RedirectResponse, PlainTextResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.exceptions import RequestValidationError
from pydantic import BaseModel, Field

from app.bot import build_dispatcher
from app.billing import (
    BillingError,
    activate_subscription,
    best_effort_ton_verify,
    build_ton_transaction,
    build_ton_payment_link,
    create_btcpay_invoice,
    PLAN_DESCRIPTIONS,
    create_telegram_stars_invoice,
    create_yookassa_payment,
    create_yookassa_credit_pack_payment,
    create_yookassa_autopayment,
    enabled_providers,
    fetch_yookassa_payment,
    make_order_id,
    normalize_plan_key,
    plan_catalog,
    price_for_provider,
    provider_enabled,
    public_plan_catalog,
    resolve_payment_provider,
    verify_btcpay_signature,
    verify_yookassa_payment,
)
from app.config import Settings, get_settings, require_runtime_settings, runtime_setting_issues
from app.ai_quality import detect_chat_intent
from app.credits import estimate_credits, estimate_output_tokens
from app.economics import credit_pack_margin, guard_credits_for_margin, plan_economics, plan_features
from app.db import CreditLimitError, Database
from app.openrouter_client import AIClientError, OpenRouterClient
from app.prompts import MODES, get_mini_app_tools
from app.rate_limit import RateLimitError, RateLimiter
from app.telegram_auth import TelegramAuthError, TelegramUser, dev_user, validate_telegram_init_data
from app.secure_store import decrypt_text, encrypt_text, mask_account, mask_token, only_digits
from app.features import FeatureStore, init_features

logger = logging.getLogger(__name__)
BASE_DIR = Path(__file__).resolve().parents[1]
STATIC_DIR = BASE_DIR / "static"


class AskRequest(BaseModel):
    mode: str = Field(default="strategy", min_length=1, max_length=64)
    text: str = Field(min_length=3, max_length=8000)


class AskResponse(BaseModel):
    answer: str
    mode: str
    saved_id: int


class GenerateRequest(BaseModel):
    tool_id: str = Field(min_length=1, max_length=64)
    user_input: str = Field(min_length=3, max_length=8000)
    optional_fields: dict[str, Any] = Field(default_factory=dict)
    telegram_user_id: int | None = None


class ToolRunRequest(BaseModel):
    tool_id: str = Field(min_length=1, max_length=64)
    input: dict[str, Any] = Field(default_factory=dict)
    telegram_user_id: int | None = None


class GenerateResponse(BaseModel):
    ok: bool
    result: str | None = None
    tool_id: str | None = None
    saved_id: int | None = None
    usage: dict[str, Any] | None = None
    error: str | None = None


class ChatRequest(BaseModel):
    telegram_user_id: int | None = None
    message: str = Field(min_length=1, max_length=8000)
    conversation_id: str | None = Field(default=None, max_length=64)


class ChatResponse(BaseModel):
    ok: bool
    answer: str | None = None
    conversation_id: str | None = None
    usage: dict[str, Any] | None = None
    error: str | None = None


class HistoryResponse(BaseModel):
    items: list[dict[str, Any]]
    conversations: list[dict[str, Any]] = Field(default_factory=list)
    tool_runs: list[dict[str, Any]] = Field(default_factory=list)
    saved: list[dict[str, Any]] = Field(default_factory=list)


class SaveResultRequest(BaseModel):
    source_type: str = Field(default="manual", min_length=1, max_length=40)
    source_id: str = Field(default="frontend", min_length=1, max_length=120)
    title: str | None = Field(default=None, max_length=160)
    content: str = Field(min_length=1, max_length=20000)


class FeedbackRequest(BaseModel):
    source_type: str | None = Field(default=None, max_length=40)
    source_id: str | None = Field(default=None, max_length=120)
    rating: int = Field(default=0, ge=-1, le=1)
    message: str | None = Field(default=None, max_length=2000)


class BillingOrderRequest(BaseModel):
    plan: str = Field(min_length=2, max_length=40)
    provider: str | None = Field(default="auto", max_length=40)
    telegram_user_id: int | None = None
    auto_renew: bool = False


class PayoutMethodRequest(BaseModel):
    bik: str = Field(min_length=9, max_length=32)
    account_number: str = Field(min_length=20, max_length=40)
    bank_name: str | None = Field(default=None, max_length=160)
    holder_name: str | None = Field(default=None, max_length=160)


class AutopaySettingsRequest(BaseModel):
    enabled: bool
    plan: str | None = Field(default=None, min_length=2, max_length=40)
    provider: str = Field(default="yookassa", max_length=40)


class TonVerifyRequest(BaseModel):
    order_id: str = Field(min_length=4, max_length=80)
    telegram_user_id: int | None = None
    wallet_address: str | None = Field(default=None, max_length=120)
    tx_hash: str | None = Field(default=None, max_length=160)


class BusinessProfileRequest(BaseModel):
    inn: str | None = Field(default=None, max_length=32)
    user_type: str | None = None
    main_goal: str | None = None
    business_name: str | None = None
    niche: str | None = None
    marketplace: str | None = None
    target_audience: str | None = None
    average_price: str | None = None
    description: str | None = None
    main_problem: str | None = None




class ProjectRequest(BaseModel):
    name: str | None = Field(default=None, max_length=120)
    business_name: str | None = Field(default=None, max_length=120)
    niche: str | None = Field(default=None, max_length=160)
    marketplace: str | None = Field(default=None, max_length=80)
    target_audience: str | None = Field(default=None, max_length=240)
    description: str | None = Field(default=None, max_length=3000)
    tone: str | None = Field(default=None, max_length=120)
    is_active: bool | None = None


class MemoryRequest(BaseModel):
    project_id: str | None = Field(default=None, max_length=80)
    category: str | None = Field(default="general", max_length=80)
    key: str | None = Field(default=None, max_length=120)
    value: str | None = Field(default=None, max_length=4000)
    content: str | None = Field(default=None, max_length=4000)
    confidence: float | None = Field(default=1.0, ge=0, le=1)
    source: str | None = Field(default="manual", max_length=80)


class TemplateRequest(BaseModel):
    title: str = Field(min_length=1, max_length=120)
    category: str | None = Field(default="general", max_length=80)
    content: str = Field(min_length=1, max_length=10000)


class CreditPackOrderRequest(BaseModel):
    pack_key: str | None = Field(default=None, min_length=3, max_length=80)
    pack_id: str | None = Field(default=None, min_length=3, max_length=80)
    provider: str | None = Field(default=None, max_length=80)
    telegram_user_id: int | None = None


class NotificationPrefsRequest(BaseModel):
    low_credits: bool | None = None
    subscription_reminders: bool | None = None
    product_updates: bool | None = None
    weekly_digest: bool | None = None


class SupportTicketRequest(BaseModel):
    subject: str | None = Field(default=None, max_length=160)
    message: str = Field(min_length=3, max_length=5000)
    category: str = Field(default="bug", max_length=40)


class SupportMessageRequest(BaseModel):
    message: str = Field(min_length=1, max_length=5000)


class SupportTicketStatusRequest(BaseModel):
    status: str = Field(default="closed", max_length=40)


class DocumentRequest(BaseModel):
    project_id: str | None = Field(default=None, max_length=80)
    title: str = Field(min_length=1, max_length=160)
    document_type: str = Field(default="note", max_length=60)
    content: str = Field(min_length=1, max_length=30000)
    source: str | None = Field(default="manual", max_length=60)
    metadata: dict[str, Any] = Field(default_factory=dict)


class DocumentGenerateRequest(BaseModel):
    project_id: str | None = Field(default=None, max_length=80)
    document_type: str = Field(default="business_plan", max_length=60)
    title: str | None = Field(default=None, max_length=160)
    prompt: str | None = Field(default=None, max_length=8000)
    mode: str | None = Field(default=None, max_length=64)


class RoadmapRequest(BaseModel):
    project_id: str | None = Field(default=None, max_length=80)
    title: str | None = Field(default=None, max_length=160)
    horizon: str | None = Field(default="30 дней", max_length=80)
    summary: str | None = Field(default=None, max_length=2000)
    tasks: list[dict[str, Any]] = Field(default_factory=list)


class TaskRequest(BaseModel):
    id: str | None = Field(default=None, max_length=80)
    roadmap_id: str | None = Field(default=None, max_length=80)
    project_id: str | None = Field(default=None, max_length=80)
    title: str = Field(min_length=1, max_length=200)
    description: str | None = Field(default=None, max_length=1000)
    status: str | None = Field(default="todo", max_length=30)
    priority: int | None = Field(default=2, ge=1, le=5)
    due_at: str | None = Field(default=None, max_length=80)


class AiFeedbackRequest(BaseModel):
    conversation_id: str | None = Field(default=None, max_length=80)
    message_id: str | None = Field(default=None, max_length=80)
    source_type: str | None = Field(default="chat", max_length=40)
    rating: int = Field(default=0, ge=-1, le=1)
    reason: str | None = Field(default=None, max_length=120)
    comment: str | None = Field(default=None, max_length=1000)


class OrganizationRequest(BaseModel):
    title: str = Field(min_length=1, max_length=120)
    description: str | None = Field(default=None, max_length=1000)


class OrganizationInviteRequest(BaseModel):
    username: str = Field(min_length=2, max_length=64)


class OrganizationAcceptRequest(BaseModel):
    token: str = Field(min_length=8, max_length=80)

class StatsResponse(BaseModel):
    used_today: int
    daily_limit: int | None
    per_minute_limit: int
    free_limit: int
    monthly_limit: int
    used_total: int
    used_period: int
    remaining: int | None
    plan: str
    status: str
    status_label: str
    subscription_until: str | None = None
    unlimited: bool = False


def _parse_money(value: Any) -> float:
    if value is None:
        return 0.0
    text = str(value).strip().replace(",", ".")
    if not text:
        return 0.0
    match = re.search(r"-?\d+(?:\.\d+)?", text.replace(" ", ""))
    return float(match.group(0)) if match else 0.0


def _parse_rate_or_amount(value: Any, base: float) -> tuple[float, float]:
    text = str(value or "").strip().replace(",", ".")
    if not text:
        return 0.0, 0.0
    raw = _parse_money(text)
    if "%" in text:
        rate = raw / 100
        return base * rate, rate
    if 0 < raw < 1:
        return base * raw, raw
    return raw, raw / base if base else 0.0


def margin_calculation(input_data: dict[str, Any]) -> dict[str, Any] | None:
    sale_price = _parse_money(input_data.get("sale_price"))
    purchase_price = _parse_money(input_data.get("purchase_price"))
    if sale_price <= 0 or purchase_price < 0:
        return None

    commission_amount, commission_rate = _parse_rate_or_amount(input_data.get("commission"), sale_price)
    taxes_amount, taxes_rate = _parse_rate_or_amount(input_data.get("taxes_other"), sale_price)
    logistics = _parse_money(input_data.get("logistics"))
    packaging = _parse_money(input_data.get("packaging"))
    ads = _parse_money(input_data.get("ads"))
    fixed_costs = purchase_price + logistics + packaging + ads
    variable_costs = commission_amount + taxes_amount
    total_costs = fixed_costs + variable_costs
    profit = sale_price - total_costs
    margin = profit / sale_price * 100 if sale_price else 0
    roi_base = fixed_costs + variable_costs
    roi = profit / roi_base * 100 if roi_base else 0

    variable_rate = commission_rate + taxes_rate
    breakeven_price = None
    if variable_rate < 1:
        breakeven_price = fixed_costs / (1 - variable_rate)

    return {
        "sale_price": sale_price,
        "purchase_price": purchase_price,
        "commission_amount": commission_amount,
        "taxes_amount": taxes_amount,
        "logistics": logistics,
        "packaging": packaging,
        "ads": ads,
        "total_costs": total_costs,
        "profit": profit,
        "margin": margin,
        "roi": roi,
        "breakeven_price": breakeven_price,
    }



def json_loads_safe(raw_body: bytes) -> dict[str, Any]:
    try:
        value = json.loads(raw_body.decode("utf-8"))
        return value if isinstance(value, dict) else {}
    except Exception:
        return {}


def _money(value: float | None) -> str:
    if value is None:
        return "недостаточно данных"
    return f"{value:.2f} руб."


def margin_markdown(calculation: dict[str, Any] | None) -> str:
    if not calculation:
        return ""
    breakeven = calculation.get("breakeven_price")
    return "\n".join(
        [
            "## Расчет backend",
            f"- Прибыль с единицы: {_money(calculation['profit'])}",
            f"- Маржа: {calculation['margin']:.1f}%",
            f"- ROI: {calculation['roi']:.1f}%",
            f"- Себестоимость с расходами: {_money(calculation['total_costs'])}",
            f"- Ориентир безубыточной цены: {_money(breakeven) if breakeven else 'недостаточно данных'}",
            "",
        ]
    )


def create_app(settings: Settings | None = None) -> FastAPI:
    settings = settings or get_settings()
    db = Database(settings.database_dsn)
    ai_client = OpenRouterClient(settings)
    rate_limiter = RateLimiter(settings, db)
    features = FeatureStore(settings.database_dsn)

    app = FastAPI(
        title="FounderPilot AI",
        description="Telegram Mini App Bot with AI business advisor",
        version="1.0.0",
    )

    app.add_middleware(
        CORSMiddleware,
        allow_origins=settings.cors_allowed_origins,
        allow_credentials=False,
        allow_methods=["GET", "POST", "OPTIONS"],
        allow_headers=["*"],
    )

    def client_ip(request: Request) -> str:
        if settings.trust_proxy_headers:
            forwarded = request.headers.get("x-forwarded-for")
            if forwarded:
                return forwarded.split(",", 1)[0].strip()[:64] or "unknown"
            real_ip = request.headers.get("x-real-ip")
            if real_ip:
                return real_ip.strip()[:64] or "unknown"
        return request.client.host if request.client else "unknown"

    def fingerprint_hash(value: str | None) -> str | None:
        clean = str(value or "").strip()
        if not clean:
            return None
        return hashlib.sha256(f"{settings.app_secret}:fingerprint:{clean}".encode("utf-8")).hexdigest()

    def add_security_headers(response):
        response.headers.setdefault("X-Content-Type-Options", "nosniff")
        response.headers.setdefault("Referrer-Policy", "strict-origin-when-cross-origin")
        response.headers.setdefault("Permissions-Policy", "camera=(), microphone=(), geolocation=()")
        response.headers.setdefault("X-Frame-Options", "SAMEORIGIN")
        response.headers.setdefault("Cross-Origin-Resource-Policy", "same-origin")
        if settings.is_public_deployment:
            response.headers.setdefault("Strict-Transport-Security", "max-age=31536000; includeSubDomains")
        return response

    @app.middleware("http")
    async def anti_abuse_middleware(request: Request, call_next):
        content_length = request.headers.get("content-length")
        if content_length:
            try:
                if int(content_length) > settings.max_request_body_bytes:
                    return add_security_headers(
                        JSONResponse(status_code=413, content={"ok": False, "error": "Request body is too large."})
                    )
            except ValueError:
                return add_security_headers(
                    JSONResponse(status_code=400, content={"ok": False, "error": "Invalid Content-Length header."})
                )

        protected = request.url.path in {
            "/api/chat",
            "/api/ask",
            "/api/generate",
            "/api/tools/run",
            "/api/billing/create-order",
            "/api/billing/ton/verify",
            "/api/credits/packs/order",
            "/api/organizations/invite",
        } or request.url.path.endswith("/invites")
        if protected:
            ip = client_ip(request)
            fingerprint = request.headers.get("x-device-fingerprint") or request.headers.get("x-fingerprint")
            recent = await features.ip_event_count(ip, settings.app_secret, minutes=10)
            risk = 25 if recent > 30 else 10 if recent > 12 else 0
            await features.record_abuse_event(
                telegram_id=None,
                ip=ip,
                fingerprint_hash=fingerprint_hash(fingerprint),
                path=request.url.path,
                event_type="api_request",
                risk_score=risk,
                metadata={"recent_ip_events_10m": recent, "user_agent": request.headers.get("user-agent", "")[:180]},
                secret=settings.app_secret,
            )
            if recent > 120:
                response = JSONResponse(status_code=429, content={"ok": False, "error": "Too many requests from this connection. Try again later."})
                return add_security_headers(response)
        return add_security_headers(await call_next(request))

    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")

    @app.exception_handler(HTTPException)
    async def http_exception_handler(request: Request, exc: HTTPException) -> JSONResponse:
        return JSONResponse(status_code=exc.status_code, content={"ok": False, "error": str(exc.detail)})

    @app.exception_handler(RequestValidationError)
    async def validation_exception_handler(request: Request, exc: RequestValidationError) -> JSONResponse:
        return JSONResponse(status_code=422, content={"ok": False, "error": "Проверьте заполнение полей и повторите запрос."})

    def resolve_telegram_id(user: TelegramUser, requested_id: int | None = None) -> int:
        if settings.allow_dev_auth and requested_id:
            return requested_id
        return user.id

    def usage_payload(access: dict[str, Any]) -> dict[str, Any]:
        return {
            "daily_used": access["used_today"],
            "daily_limit": access["current_limit"],
            "used_total": access["used_total"],
            "used_period": access["used_period"],
            "remaining": access["remaining"],
            "plan": access["plan"],
            "status": access["status"],
            "status_label": access["status_label"],
            "subscription_until": access["subscription_until"],
            "unlimited": access["unlimited"],
            "unit_name": "кредиты",
            "credits_daily_limit": access.get("daily_limit"),
            "credits_monthly_limit": access.get("monthly_limit"),
            "credits_used_today": access.get("credits_used_today", access.get("used_today", 0)),
            "credits_used_month": access.get("credits_used_period", access.get("used_period", 0)),
            "credits_remaining": access.get("remaining"),
            "credits_reserved": access.get("credits_reserved", 0),
            "daily_remaining": access.get("daily_remaining"),
            "monthly_remaining": access.get("monthly_remaining"),
        }

    def stored_user_input(user_input: str, optional_fields: dict[str, Any]) -> str:
        filled = [
            (key, str(value).strip())
            for key, value in optional_fields.items()
            if value is not None and str(value).strip()
        ]
        if not filled:
            return user_input.strip()
        fields = "\n".join(f"{key}: {value}" for key, value in filled)
        return f"{user_input.strip()}\n\nДополнительные поля:\n{fields}"

    def prepare_tool_context(tool_id: str, input_data: dict[str, Any]) -> tuple[str, dict[str, Any], str]:
        optional_fields = dict(input_data)
        input_text = "\n".join(
            f"{key}: {value}" for key, value in input_data.items() if value is not None and str(value).strip()
        )
        prefix = ""
        if tool_id == "margin_calc":
            calculation = margin_calculation(input_data)
            if calculation:
                prefix = margin_markdown(calculation)
                optional_fields["backend_margin_calculation"] = prefix
                input_text = f"{input_text}\n\nСерверный расчет:\n{prefix}".strip()
        return input_text, optional_fields, prefix

    def parse_conversation_id(value: str | None) -> str | None:
        if not value:
            return None
        clean = value.strip()
        return clean or None

    def business_context(profile: dict[str, Any] | None) -> str:
        if not profile:
            return ""
        labels = {
            "user_type": "Тип пользователя",
            "main_goal": "Главная цель",
            "business_name": "Бизнес",
            "niche": "Ниша",
            "marketplace": "Маркетплейс",
            "target_audience": "Целевая аудитория",
            "average_price": "Средняя цена",
            "description": "Описание",
            "main_problem": "Главная проблема",
        }
        lines = [f"{label}: {profile.get(key)}" for key, label in labels.items() if profile.get(key)]
        return "\n".join(lines)

    def ai_user_context(user: TelegramUser, profile: dict[str, Any] | None = None, organizations: dict[str, Any] | None = None) -> str:
        lines: list[str] = []
        if user.first_name:
            lines.append(f"Имя пользователя в Telegram: {user.first_name}")
        if user.username:
            lines.append(f"Username Telegram: @{user.username}")
        if profile and profile.get("plan"):
            plan_key = str(profile.get("plan") or "free").lower()
            lines.append(f"Активный тариф FounderPilot: {plan_key}")
            if profile.get("subscription_until") or profile.get("subscription_expires_at"):
                lines.append(f"Подписка действует до: {profile.get('subscription_until') or profile.get('subscription_expires_at')}")
            if profile.get("daily_limit"):
                lines.append(f"Дневной лимит тарифа: {profile.get('daily_limit')} кредитов")
            if profile.get("monthly_limit"):
                lines.append(f"Месячный лимит тарифа: {profile.get('monthly_limit')} кредитов")
            if plan_key in {"pro", "business"}:
                lines.append("Отвечай с уровнем глубины платного тарифа: больше конкретики, шагов, рисков и чисел.")
        active_org = (organizations or {}).get("active") if isinstance(organizations, dict) else None
        if active_org:
            org_title = active_org.get("title") or active_org.get("organization_title") or active_org.get("name")
            org_role = active_org.get("role") or active_org.get("member_role")
            if org_title:
                lines.append(f"Активная организация: {org_title}")
            if org_role:
                lines.append(f"Роль в организации: {org_role}")
        return "\n".join(lines)

    def subscription_next_charge_at() -> str:
        return (datetime.now(timezone.utc) + timedelta(days=30)).isoformat()

    def validate_payout_payload(payload: PayoutMethodRequest) -> tuple[str, str]:
        bik = only_digits(payload.bik)
        account = only_digits(payload.account_number)
        if len(bik) != 9:
            raise HTTPException(status_code=422, detail="БИК должен состоять из 9 цифр.")
        if len(account) != 20:
            raise HTTPException(status_code=422, detail="Номер счёта должен состоять из 20 цифр. Не указывайте номер карты.")
        return bik, account

    def metadata_from_order(order: dict[str, Any] | None) -> dict[str, Any]:
        if not order:
            return {}
        raw = order.get("metadata_json")
        if not raw:
            return {}
        try:
            data = json.loads(raw)
        except (TypeError, ValueError):
            return {}
        return data if isinstance(data, dict) else {}

    async def finalize_yookassa_order(order: dict[str, Any], raw_event: dict[str, Any] | None = None) -> dict[str, Any]:
        """Verify YooKassa payment and activate the subscription exactly once.

        This makes sales reliable even if the user returns to the Mini App before
        the webhook arrives. Polling /api/billing/order can finish activation too.
        """
        if not order:
            raise BillingError("Заказ не найден.")
        if str(order.get("status") or "").lower() == "paid":
            return order
        external_id = str(order.get("external_payment_id") or "")
        if not external_id:
            raise BillingError("У заказа нет ID платежа ЮKassa.")
        plans = plan_catalog(settings)
        plan = plans.get(str(order.get("plan") or "").lower())
        if not plan or plan.key == "free":
            raise BillingError("Некорректный тариф в заказе.")

        verified_payment = await verify_yookassa_payment(settings, order, external_id)
        await activate_subscription(db, order["telegram_user_id"], plan.key, "yookassa", order["id"], plan.daily_limit, plan.monthly_limit)
        await notify_user(order["telegram_user_id"], title="Подписка активна", body=f"Тариф {plan.title} успешно подключён.", type="billing", action_url="subscription", metadata={"order_id": order["id"], "plan": plan.key})

        method = verified_payment.get("payment_method") or {}
        method_id = method.get("id")
        saved_meta = metadata_from_order(order)
        payment_meta = verified_payment.get("metadata") or {}
        card = method.get("card") or {}
        last4 = card.get("last4") or card.get("last_4") or card.get("last_four")
        method_title = method.get("title") or method.get("saved_payment_method_title")
        if last4:
            card_type = str(card.get("card_type") or card.get("issuer_name") or "Карта").strip() or "Карта"
            payment_method_mask = f"{card_type} •••• {last4}"
        elif method_title:
            payment_method_mask = str(method_title)
        else:
            payment_method_mask = f"Карта {mask_token(str(method_id))}"
        if settings.yookassa_enable_saved_payment_method and method_id and (saved_meta.get("auto_renew") or str(payment_meta.get("auto_renew")) == "1"):
            await db.save_autopay_payment_method(
                int(order["telegram_user_id"]),
                plan=plan.key,
                payment_method_id_encrypted=encrypt_text(str(method_id), settings.app_secret) or "",
                payment_method_mask=payment_method_mask,
                next_charge_at=subscription_next_charge_at(),
            )

        await db.record_payment(
            order_id=order["id"],
            telegram_id=int(order["telegram_user_id"]),
            provider="yookassa",
            plan=plan.key,
            amount=float(order["amount"]),
            currency=order["currency"],
            status="succeeded",
            external_payment_id=external_id,
            raw_event=raw_event or verified_payment,
        )
        updated = await db.get_billing_order(order["id"])
        return updated or order

    async def current_user(x_telegram_init_data: str | None = Header(default=None)) -> TelegramUser:
        if settings.allow_dev_auth:
            return dev_user()
        try:
            return validate_telegram_init_data(
                x_telegram_init_data or "",
                settings.bot_token,
                max_age_seconds=settings.telegram_init_data_max_age_seconds,
            )
        except TelegramAuthError as exc:
            raise HTTPException(status_code=401, detail=str(exc)) from exc

    def require_admin(x_admin_secret: str | None) -> None:
        if not settings.admin_secret or not x_admin_secret or not hmac.compare_digest(x_admin_secret, settings.admin_secret):
            raise HTTPException(status_code=403, detail="Доступ запрещён.")

    async def apply_profit_guard_to_estimate(telegram_id: int, estimate):
        """Raise request price in credits when the configured model/provider gets expensive.

        The guard is server-side only: the frontend may show any nice numbers it wants,
        but the backend always calculates the safe credit charge before reserving limits.
        """
        if not settings.profit_guard_enabled:
            return estimate
        access = await db.get_access_state(
            telegram_id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        plans = plan_catalog(settings)
        plan_key = str(access.get("raw_plan") or access.get("plan") or "free").lower()
        plan = plans.get(plan_key) or plans.get("free")
        guarded = guard_credits_for_margin(settings, plan, estimate)
        if guarded.credits <= estimate.credits:
            return estimate
        return replace(
            estimate,
            credits=guarded.credits,
            reason=f"{estimate.reason}; {guarded.reason}",
        )

    async def read_json_dict(request: Request) -> dict[str, Any]:
        raw_body = await request.body()
        if len(raw_body) > settings.max_request_body_bytes:
            raise HTTPException(status_code=413, detail="Request body is too large.")
        try:
            value = json.loads(raw_body.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise HTTPException(status_code=400, detail="Invalid JSON body.") from exc
        if not isinstance(value, dict):
            raise HTTPException(status_code=400, detail="JSON body must be an object.")
        return value

    async def notify_user(telegram_id: int | str, *, title: str, body: str = "", type: str = "system", action_url: str | None = None, metadata: dict[str, Any] | None = None) -> None:
        try:
            await features.create_notification(telegram_id, title=title, body=body, type=type, action_url=action_url, metadata=metadata or {})
        except Exception as exc:  # noqa: BLE001
            with suppress(Exception):
                await db.log_error("notification_create", str(exc), int(telegram_id) if str(telegram_id).isdigit() else None)

    def support_ticket_public_text(ticket: dict[str, Any], message: str) -> str:
        username = ticket.get("username") or ""
        user_name = ticket.get("user_name") or "Пользователь"
        user_line = f"@{username}" if username else str(ticket.get("telegram_user_id") or "—")
        return (
            f"<b>Новый тикет FounderPilot</b>\n"
            f"ID: <code>{html_escape(str(ticket.get('id') or ''))}</code>\n"
            f"Пользователь: <b>{html_escape(user_name)}</b> ({html_escape(user_line)})\n"
            f"Тариф: <code>{html_escape(str(ticket.get('plan') or 'free'))}</code>\n"
            f"Категория: <code>{html_escape(str(ticket.get('category') or 'bug'))}</code>\n"
            f"Тема: <b>{html_escape(str(ticket.get('subject') or 'Обращение'))}</b>\n\n"
            f"{html_escape(message)}\n\n"
            f"<i>Ответьте именно reply на это сообщение — ответ автоматически появится у пользователя в Mini App.</i>"
        )

    async def send_support_ticket_to_group(ticket: dict[str, Any], message: str) -> bool:
        if not settings.bot_token or not settings.support_group_chat_id:
            return False
        bot = Bot(token=settings.bot_token)
        try:
            kwargs: dict[str, Any] = {}
            if settings.support_group_thread_id:
                kwargs["message_thread_id"] = settings.support_group_thread_id
            sent = await bot.send_message(
                settings.support_group_chat_id,
                support_ticket_public_text(ticket, message),
                parse_mode="HTML",
                **kwargs,
            )
            await features.update_support_ticket_bridge(ticket["id"], sent.chat.id, sent.message_id)
            return True
        except Exception as exc:  # noqa: BLE001
            await db.log_error("support_group_send", str(exc), int(ticket.get("telegram_user_id") or 0) or None)
            return False
        finally:
            await bot.session.close()

    async def send_support_followup_to_group(ticket: dict[str, Any], message: str) -> bool:
        if not settings.bot_token or not settings.support_group_chat_id:
            return False
        bot = Bot(token=settings.bot_token)
        try:
            kwargs: dict[str, Any] = {}
            if settings.support_group_thread_id:
                kwargs["message_thread_id"] = settings.support_group_thread_id
            if ticket.get("group_message_id"):
                kwargs["reply_to_message_id"] = int(ticket["group_message_id"])
            sent = await bot.send_message(
                settings.support_group_chat_id,
                f"<b>Новое сообщение по тикету</b> <code>{html_escape(str(ticket.get('id') or ''))}</code>\n\n{html_escape(message)}",
                parse_mode="HTML",
                **kwargs,
            )
            await features.update_support_ticket_bridge(ticket["id"], sent.chat.id, sent.message_id)
            return True
        except Exception as exc:  # noqa: BLE001
            await db.log_error("support_group_followup", str(exc), int(ticket.get("telegram_user_id") or 0) or None)
            return False
        finally:
            await bot.session.close()

    async def send_organization_invite(invite: dict[str, Any]) -> bool:
        invited_id = invite.get("invited_telegram_user_id")
        if not invited_id or not settings.bot_token:
            return False
        url = f"{settings.webapp_url}?invite={invite.get('token')}"
        bot = Bot(token=settings.bot_token)
        try:
            await bot.send_message(
                int(invited_id),
                (
                    "<b>Приглашение в FounderPilot</b>\n\n"
                    f"Вас пригласили в организацию: <b>{invite.get('organization_title') or 'Компания'}</b>.\n"
                    "Нажмите кнопку ниже, чтобы вступить."
                ),
                parse_mode="HTML",
                reply_markup=InlineKeyboardMarkup(
                    inline_keyboard=[[InlineKeyboardButton(text="Вступить", web_app=WebAppInfo(url=url))]]
                ),
            )
            return True
        except Exception as exc:  # noqa: BLE001
            await db.log_error("organization_invite_notify", str(exc), int(invited_id))
            return False
        finally:
            await bot.session.close()

    @app.on_event("startup")
    async def on_startup() -> None:
        app.state.db_ready = False
        app.state.startup_error = None
        try:
            await db.init()
            await init_features(settings.database_dsn)
            app.state.db_ready = True
            logger.info("Database initialized at %s", settings.database_dsn)
        except Exception as exc:  # noqa: BLE001
            app.state.startup_error = str(exc)
            logger.exception("Database/features initialization failed. /health will stay alive; app APIs may fail until this is fixed: %s", exc)

    @app.get("/", include_in_schema=False)
    async def root() -> RedirectResponse:
        return RedirectResponse(url="/app")

    @app.get("/health")
    async def health() -> dict[str, Any]:
        # Railway healthcheck must stay lightweight and must not depend on Telegram,
        # OpenRouter, YooKassa or a long DB migration. Detailed warnings are in logs.
        return {
            "ok": True,
            "service": "FounderPilot AI",
            "db_ready": bool(getattr(app.state, "db_ready", False)),
        }

    @app.get("/ready")
    async def ready() -> dict[str, Any]:
        if not bool(getattr(app.state, "db_ready", False)):
            raise HTTPException(status_code=503, detail=getattr(app.state, "startup_error", None) or "Database is not ready")
        return {"ok": True, "service": "FounderPilot AI", "db_ready": True}

    @app.get("/tonconnect-manifest.json", include_in_schema=False)
    async def ton_manifest() -> dict[str, Any]:
        return {
            "url": settings.webapp_public_url,
            "name": "FounderPilot AI",
            "iconUrl": f"{settings.webapp_public_url}/static/icon-512.png",
        }

    @app.get("/app", include_in_schema=False)
    async def mini_app() -> FileResponse:
        return FileResponse(STATIC_DIR / "index.html")

    @app.get("/legal", include_in_schema=False)
    @app.get("/requisites", include_in_schema=False)
    async def legal_page() -> FileResponse:
        return FileResponse(STATIC_DIR / "legal.html")

    @app.get("/api/modes")
    async def modes() -> dict[str, Any]:
        return {
            "modes": [
                {"key": mode.key, "title": mode.title, "description": mode.description}
                for mode in MODES.values()
            ]
        }

    @app.get("/api/tools")
    async def tools() -> dict[str, Any]:
        return {
            "tools": [
                {
                    "id": tool.key,
                    "title": tool.title,
                    "description": tool.description,
                    "icon": tool.icon,
                    "placeholder": tool.placeholder,
                    "fields": list(tool.fields),
                }
                for tool in get_mini_app_tools()
            ]
        }

    @app.get("/api/me")
    async def me(start_param: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url, track_visit=True)
        await db.expire_subscription_if_needed(user.id)
        if start_param:
            await db.set_referrer(user.id, start_param.strip())
        access = await db.get_access_state(
            user.id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        profile_data = await db.get_business_profile(user.id)
        user_profile = await db.get_user_profile(user.id)
        usage = usage_payload(access)
        active_project = await features.get_active_project(user.id)
        organizations = await features.list_organizations(user.id, user.username)
        payout_method = await db.get_payout_method(user.id)
        autopay = await db.get_autopay_settings(user.id)
        stored_first = (user_profile or {}).get("first_name")
        stored_last = (user_profile or {}).get("last_name")
        stored_username = (user_profile or {}).get("username")
        stored_photo = (user_profile or {}).get("photo_url")
        return {
            "user": {
                "id": user.id,
                "telegram_id": user.id,
                "telegram_user_id": str(user.id),
                "created_at": (user_profile or {}).get("created_at"),
                "updated_at": (user_profile or {}).get("updated_at"),
                "last_seen_at": (user_profile or {}).get("last_seen_at"),
                "login_count": int((user_profile or {}).get("login_count") or 0),
                "subscription_until": (user_profile or {}).get("subscription_until") or (user_profile or {}).get("subscription_expires_at"),
                "username": user.username or stored_username,
                "first_name": user.first_name or stored_first or "Пользователь",
                "last_name": user.last_name or stored_last,
                "photo_url": user.photo_url or stored_photo,
                "avatar_url": user.photo_url or stored_photo,
                "plan": (user_profile or {}).get("plan", "free"),
                "is_admin": settings.is_admin(user.id),
                "onboarding_completed": bool((user_profile or {}).get("onboarding_completed")),
                **usage,
            },
            "profile": profile_data,
            "active_project": active_project,
            "organization": organizations.get("active"),
            "organizations": organizations,
            "payout_method": payout_method,
            "autopay": autopay,
            "usage": usage,
        }

    @app.get("/api/ai/status")
    async def ai_status(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        await db.expire_subscription_if_needed(user.id)
        user_profile = await db.get_user_profile(user.id)
        plan_key = str((user_profile or {}).get("plan") or "free").lower()
        access = await db.get_access_state(
            user.id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        return {
            "ok": True,
            "plan": plan_key,
            "model": ai_client.model_for_plan(plan_key),
            "fallback_models": settings.openrouter_fallback_models,
            "quality_mode": settings.ai_answer_quality_mode,
            "limits": usage_payload(access),
            "configured": bool(settings.openrouter_api_key.strip()),
        }

    @app.post("/api/onboarding")
    async def onboarding(payload: BusinessProfileRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        profile_data = await db.complete_onboarding(user.id, payload.model_dump())
        return {"ok": True, "profile": profile_data}

    @app.get("/api/business-profile")
    async def business_profile(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "profile": await db.get_business_profile(user.id)}

    @app.post("/api/business-profile")
    async def save_business_profile(payload: BusinessProfileRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        profile_data = await db.upsert_business_profile(user.id, payload.model_dump())
        return {"ok": True, "profile": profile_data}

    @app.post("/api/profile/save")
    async def profile_save(request: Request, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        raw = await read_json_dict(request)
        description = raw.get("business_profile") or raw.get("description") or raw.get("text")
        company_name = raw.get("company_name") or raw.get("company") or raw.get("business_name")
        inn = raw.get("inn")
        payload = {"description": description, "business_name": company_name, "inn": inn}
        profile_data = await db.upsert_business_profile(user.id, payload)
        return {"ok": True, "profile": profile_data}

    @app.delete("/api/business-profile")
    async def delete_business_profile(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        await db.clear_business_profile(user.id)
        return {"ok": True}

    @app.post("/api/ask", response_model=AskResponse)
    async def ask(payload: AskRequest, user: TelegramUser = Depends(current_user)) -> AskResponse:
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        await db.expire_subscription_if_needed(user.id)
        user_profile = await db.get_user_profile(user.id)
        plan_key = str((user_profile or {}).get("plan") or "free").lower()
        ai_model = ai_client.model_for_plan(plan_key)
        access_state = await db.get_access_state(
            user.id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        estimate = estimate_credits(payload.mode, payload.text, model=ai_model)
        estimate = await apply_profit_guard_to_estimate(user.id, estimate)

        try:
            await rate_limiter.check(user.id, estimate.credits)
            await db.reserve_credits(
                user.id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
                metadata={"endpoint": "/api/ask", "reason": estimate.reason, "model": ai_model},
            )
            ask_optional: dict[str, Any] = {}
            profile_data = await db.get_business_profile(user.id)
            organizations = await features.list_organizations(user.id, user.username)
            user_context = ai_user_context(user, user_profile, organizations)
            context = business_context(profile_data)
            project_context = await features.project_context_text(user.id)
            if user_context:
                ask_optional["telegram_user_context"] = user_context
            if context:
                ask_optional["business_profile"] = context
            if project_context:
                ask_optional["project_memory"] = project_context
            answer = await ai_client.ask_business_ai(
                payload.mode,
                payload.text,
                ask_optional,
                plan_key=plan_key,
                access_context=access_state,
            )
            output_tokens = estimate_output_tokens(answer)
            await db.finalize_credit_charge(
                user.id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                estimate.credits,
                model=ai_model,
                input_tokens=estimate.input_tokens_estimated,
                output_tokens=output_tokens,
            )
            saved_id = await db.save_request(user.id, payload.mode, payload.text, answer)
        except (RateLimitError, CreditLimitError) as exc:
            raise HTTPException(status_code=429, detail=str(exc)) from exc
        except AIClientError as exc:
            await db.refund_reserved_credits(user.id, estimate.request_id, str(exc), tool_id=estimate.tool_id, estimated_credits=estimate.credits, model=ai_model)
            raise HTTPException(status_code=502, detail=str(exc)) from exc
        except Exception as exc:  # noqa: BLE001
            await db.refund_reserved_credits(user.id, estimate.request_id, str(exc), tool_id=estimate.tool_id, estimated_credits=estimate.credits, model=ai_model)
            raise

        return AskResponse(answer=answer, mode=payload.mode, saved_id=saved_id)

    @app.post("/api/generate", response_model=GenerateResponse)
    async def generate(payload: GenerateRequest, user: TelegramUser = Depends(current_user)) -> GenerateResponse:
        telegram_id = resolve_telegram_id(user, payload.telegram_user_id)
        await db.upsert_user(telegram_id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        await db.expire_subscription_if_needed(telegram_id)

        if payload.tool_id not in MODES:
            return GenerateResponse(ok=False, error="Неизвестный AI-инструмент. Обновите приложение и попробуйте снова.")

        estimate = None
        user_profile = await db.get_user_profile(telegram_id)
        plan_key = str((user_profile or {}).get("plan") or "free").lower()
        ai_model = ai_client.model_for_plan(plan_key)
        try:
            profile_data = await db.get_business_profile(telegram_id)
            raw_fields = dict(payload.optional_fields)
            raw_fields["user_input"] = payload.user_input
            input_text, optional_fields, result_prefix = prepare_tool_context(payload.tool_id, raw_fields)
            context = business_context(profile_data)
            project_context = await features.project_context_text(telegram_id)
            access_state = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            user_context = ai_user_context(user, user_profile)
            if user_context:
                optional_fields["telegram_user_context"] = user_context
            if context:
                optional_fields["business_profile"] = context
            if project_context:
                optional_fields["project_memory"] = project_context
            estimate = estimate_credits(
                payload.tool_id,
                input_text or payload.user_input,
                model=ai_model,
                optional_fields=optional_fields,
            )
            estimate = await apply_profit_guard_to_estimate(telegram_id, estimate)
            await rate_limiter.check(telegram_id, estimate.credits)
            await db.reserve_credits(
                telegram_id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
                metadata={"endpoint": "/api/generate", "reason": estimate.reason, "model": ai_model},
            )
            ai_answer = await ai_client.ask_business_ai(
                payload.tool_id,
                input_text or payload.user_input,
                optional_fields,
                plan_key=plan_key,
                access_context=access_state,
            )
            answer = f"{result_prefix}\n{ai_answer}".strip() if result_prefix else ai_answer
            await db.finalize_credit_charge(
                telegram_id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                estimate.credits,
                model=ai_model,
                input_tokens=estimate.input_tokens_estimated,
                output_tokens=estimate_output_tokens(answer),
            )
            saved_id = await db.save_request(
                telegram_id,
                payload.tool_id,
                stored_user_input(payload.user_input, optional_fields),
                answer,
            )
            await db.create_tool_run(
                telegram_id,
                payload.tool_id,
                {"user_input": payload.user_input, "optional_fields": payload.optional_fields},
                result_text=answer,
                model=ai_model,
                tokens_used=estimate.input_tokens_estimated + estimate_output_tokens(answer),
            )
            access = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
        except (RateLimitError, CreditLimitError) as exc:
            access = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            return GenerateResponse(ok=False, error=str(exc), tool_id=payload.tool_id, usage=usage_payload(access))
        except AIClientError as exc:
            if estimate is not None:
                await db.refund_reserved_credits(
                    telegram_id,
                    estimate.request_id,
                    str(exc),
                    tool_id=estimate.tool_id,
                    estimated_credits=estimate.credits,
                    model=ai_model,
                )
            await db.log_error("generate", str(exc), telegram_id)
            return GenerateResponse(ok=False, error=str(exc), tool_id=payload.tool_id)
        except Exception as exc:  # noqa: BLE001
            if estimate is not None:
                await db.refund_reserved_credits(
                    telegram_id,
                    estimate.request_id,
                    str(exc),
                    tool_id=estimate.tool_id,
                    estimated_credits=estimate.credits,
                    model=ai_model,
                )
            logger.exception("Mini App generation failed")
            await db.log_error("generate", str(exc), telegram_id)
            return GenerateResponse(ok=False, error=f"Непредвиденная ошибка: {exc}", tool_id=payload.tool_id)

        return GenerateResponse(
            ok=True,
            result=answer,
            tool_id=payload.tool_id,
            saved_id=saved_id,
            usage=usage_payload(access),
        )

    @app.post("/api/tools/run")
    async def tools_run(payload: ToolRunRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        telegram_id = resolve_telegram_id(user, payload.telegram_user_id)
        await db.expire_subscription_if_needed(telegram_id)
        if payload.tool_id not in MODES:
            return {"ok": False, "error": "Неизвестный AI-инструмент."}

        input_text, optional_fields, result_prefix = prepare_tool_context(payload.tool_id, payload.input)
        if len(input_text.strip()) < 3:
            return {"ok": False, "error": "Заполните хотя бы одно поле инструмента."}

        estimate = None
        user_profile = await db.get_user_profile(telegram_id)
        plan_key = str((user_profile or {}).get("plan") or "free").lower()
        ai_model = ai_client.model_for_plan(plan_key)
        try:
            profile_data = await db.get_business_profile(telegram_id)
            context = business_context(profile_data)
            project_context = await features.project_context_text(telegram_id)
            access_state = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            user_context = ai_user_context(user, user_profile)
            if user_context:
                optional_fields["telegram_user_context"] = user_context
            if context:
                optional_fields["business_profile"] = context
            if project_context:
                optional_fields["project_memory"] = project_context
            estimate = estimate_credits(
                payload.tool_id,
                input_text,
                model=ai_model,
                optional_fields=optional_fields,
            )
            estimate = await apply_profit_guard_to_estimate(telegram_id, estimate)
            await rate_limiter.check(telegram_id, estimate.credits)
            await db.reserve_credits(
                telegram_id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
                metadata={"endpoint": "/api/tools/run", "reason": estimate.reason, "model": ai_model},
            )
            ai_answer = await ai_client.ask_business_ai(
                payload.tool_id,
                input_text,
                optional_fields,
                plan_key=plan_key,
                access_context=access_state,
            )
            answer = f"{result_prefix}\n{ai_answer}".strip() if result_prefix else ai_answer
            await db.finalize_credit_charge(
                telegram_id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                estimate.credits,
                model=ai_model,
                input_tokens=estimate.input_tokens_estimated,
                output_tokens=estimate_output_tokens(answer),
            )
            tool_run_id = await db.create_tool_run(
                telegram_id,
                payload.tool_id,
                payload.input,
                result_text=answer,
                model=ai_model,
                tokens_used=estimate.input_tokens_estimated + estimate_output_tokens(answer),
            )
            await db.save_request(telegram_id, payload.tool_id, input_text, answer)
            access = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            return {"ok": True, "result": answer, "tool_run_id": str(tool_run_id), "usage": usage_payload(access)}
        except (RateLimitError, CreditLimitError) as exc:
            access = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            return {"ok": False, "error": str(exc), "usage": usage_payload(access)}
        except AIClientError as exc:
            if estimate is not None:
                await db.refund_reserved_credits(telegram_id, estimate.request_id, str(exc), tool_id=estimate.tool_id, estimated_credits=estimate.credits, model=ai_model)
            await db.create_tool_run(telegram_id, payload.tool_id, payload.input, status="error", error_message=str(exc))
            await db.log_error("tools_run", str(exc), telegram_id)
            return {"ok": False, "error": str(exc)}
        except Exception as exc:  # noqa: BLE001
            if estimate is not None:
                await db.refund_reserved_credits(telegram_id, estimate.request_id, str(exc), tool_id=estimate.tool_id, estimated_credits=estimate.credits, model=ai_model)
            logger.exception("Tool run failed")
            await db.create_tool_run(telegram_id, payload.tool_id, payload.input, status="error", error_message=str(exc))
            await db.log_error("tools_run", str(exc), telegram_id)
            return {"ok": False, "error": f"Непредвиденная ошибка: {exc}"}

    @app.post("/api/chat", response_model=ChatResponse)
    async def chat(payload: ChatRequest, user: TelegramUser = Depends(current_user)) -> ChatResponse:
        telegram_id = resolve_telegram_id(user, payload.telegram_user_id)
        message = payload.message.strip()
        if len(message) < 1:
            return ChatResponse(ok=False, error="Напишите сообщение.")

        await db.upsert_user(telegram_id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        await db.expire_subscription_if_needed(telegram_id)

        estimate = None
        conversation_id = payload.conversation_id
        user_profile = await db.get_user_profile(telegram_id)
        plan_key = str((user_profile or {}).get("plan") or "free").lower()
        ai_model = ai_client.model_for_plan(plan_key)
        intent = detect_chat_intent(message)
        try:
            conversation_id = await db.get_or_create_conversation(
                telegram_id,
                parse_conversation_id(payload.conversation_id),
                message,
            )
            history = await db.list_chat_messages(conversation_id, limit=settings.ai_chat_history_messages)
            history_text = "\n".join(str(item.get("content") or "") for item in history[-settings.ai_chat_history_messages:])
            estimate = estimate_credits(
                "chat",
                message,
                history_text=history_text,
                model=ai_model,
            )
            estimate = await apply_profit_guard_to_estimate(telegram_id, estimate)
            await rate_limiter.check(telegram_id, estimate.credits)
            await db.reserve_credits(
                telegram_id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
                metadata={
                    "endpoint": "/api/chat",
                    "conversation_id": str(conversation_id),
                    "reason": estimate.reason,
                    "model": ai_model,
                    "intent": intent.key,
                },
            )
            # Save the user's message before calling the AI provider.
            # If the user closes the Mini App while the model is still answering,
            # the dialogue still remains in history instead of disappearing.
            await db.add_chat_message(conversation_id, "user", message)
            profile_data = await db.get_business_profile(telegram_id)
            organizations = await features.list_organizations(telegram_id, user.username)
            access_state = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            context_parts = [
                ai_user_context(user, user_profile, organizations),
                business_context(profile_data),
                await features.project_context_text(telegram_id),
            ]
            answer = await ai_client.ask_chat(
                history,
                message,
                "\n\n".join(part for part in context_parts if part),
                plan_key=plan_key,
                access_context=access_state,
                intent=intent,
            )
            await db.finalize_credit_charge(
                telegram_id,
                estimate.request_id,
                estimate.tool_id,
                estimate.credits,
                estimate.credits,
                model=ai_model,
                input_tokens=estimate.input_tokens_estimated,
                output_tokens=estimate_output_tokens(answer),
            )
            await db.add_chat_message(
                conversation_id,
                "assistant",
                answer,
                model=ai_model,
                tokens_used=estimate.input_tokens_estimated + estimate_output_tokens(answer),
            )
            access = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
        except (RateLimitError, CreditLimitError) as exc:
            access = await db.get_access_state(
                telegram_id,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
            )
            return ChatResponse(ok=False, error=str(exc), conversation_id=str(conversation_id) if conversation_id else payload.conversation_id, usage=usage_payload(access))
        except AIClientError as exc:
            if estimate is not None:
                await db.refund_reserved_credits(telegram_id, estimate.request_id, str(exc), tool_id=estimate.tool_id, estimated_credits=estimate.credits, model=ai_model)
            await db.log_error("chat", str(exc), telegram_id)
            return ChatResponse(ok=False, error=str(exc), conversation_id=str(conversation_id) if conversation_id else payload.conversation_id)
        except Exception as exc:  # noqa: BLE001
            if estimate is not None:
                await db.refund_reserved_credits(telegram_id, estimate.request_id, str(exc), tool_id=estimate.tool_id, estimated_credits=estimate.credits, model=ai_model)
            logger.exception("AI Chat request failed")
            await db.log_error("chat", str(exc), telegram_id)
            return ChatResponse(ok=False, error=f"Непредвиденная ошибка: {exc}", conversation_id=str(conversation_id) if conversation_id else payload.conversation_id)

        return ChatResponse(
            ok=True,
            answer=answer,
            conversation_id=str(conversation_id),
            usage=usage_payload(access),
        )

    @app.get("/api/conversations")
    async def conversations(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        items = await db.list_conversations(user.id, limit=20)
        return {"items": items}

    @app.post("/api/conversations")
    async def create_conversation(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        conversation_id = await db.create_conversation(user.id, "Новый диалог")
        return {"ok": True, "conversation_id": conversation_id}

    @app.get("/api/conversations/{conversation_id}")
    @app.get("/api/chat/{conversation_id}")
    async def chat_history(conversation_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        conversation = await db.get_conversation(user.id, conversation_id)
        if not conversation:
            raise HTTPException(status_code=404, detail="Диалог не найден.")
        messages = await db.list_chat_messages(conversation_id, limit=50)
        return {"conversation": conversation, "messages": messages}

    @app.delete("/api/conversations/{conversation_id}")
    async def delete_conversation(conversation_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        deleted = await db.archive_conversation(user.id, conversation_id)
        return {"ok": deleted}

    @app.get("/api/history", response_model=HistoryResponse)
    async def history(user: TelegramUser = Depends(current_user)) -> HistoryResponse:
        items = await db.list_recent_requests(user.id, limit=10)
        for item in items:
            mode = MODES.get(str(item.get("mode")))
            item["tool_title"] = mode.title if mode else item.get("mode")
        conversations = await db.list_conversations(user.id, limit=10)
        tool_runs = await db.list_tool_runs(user.id, limit=20)
        saved = await db.list_saved_results(user.id, limit=20)
        return HistoryResponse(items=items, conversations=conversations, tool_runs=tool_runs, saved=saved)

    @app.get("/api/saved")
    async def saved(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"items": await db.list_saved_results(user.id)}

    @app.post("/api/saved")
    async def save_result(payload: SaveResultRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        saved_id = await db.save_result(user.id, payload.source_type, payload.source_id, payload.title, payload.content)
        return {"ok": True, "id": saved_id}

    @app.delete("/api/saved/{saved_id}")
    async def delete_saved(saved_id: int, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await db.delete_saved_result(user.id, saved_id)}

    @app.post("/api/feedback")
    async def feedback(payload: FeedbackRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        feedback_id = await db.save_feedback(
            user.id,
            rating=payload.rating,
            message=payload.message,
            source_type=payload.source_type,
            source_id=payload.source_id,
        )
        return {"ok": True, "id": feedback_id}

    @app.get("/api/billing/plans")
    async def billing_plans() -> dict[str, Any]:
        providers = enabled_providers(settings)
        plans = []
        for plan in public_plan_catalog(settings).values():
            payment_providers = [] if plan.key == "free" else providers
            plans.append(
                {
                    "id": plan.key,
                    "key": plan.key,
                    "title": plan.title,
                    "name": plan.title,
                    "description": PLAN_DESCRIPTIONS.get(plan.key, ""),
                    "daily_limit": plan.daily_limit,
                    "monthly_limit": plan.monthly_limit,
                    "credits_daily_limit": plan.daily_limit,
                    "credits_monthly_limit": plan.monthly_limit,
                    "unit_name": "кредиты",
                    "price": "0 ₽" if plan.key == "free" else f"{int(plan.price_rub):,}".replace(",", " ") + " ₽ / мес",
                    "price_text": "0 ₽" if plan.key == "free" else f"{int(plan.price_rub):,}".replace(",", " ") + " ₽ / мес",
                    "price_rub": float(plan.price_rub),
                    "price_stars": plan.price_stars,
                    "price_ton": str(plan.price_ton),
                    "price_btc": str(plan.price_btc),
                    "duration_days": 30 if plan.key != "free" else None,
                    "providers": payment_providers,
                    "features": plan_features(plan.key),
                    "profit_guard_enabled": settings.profit_guard_enabled,
                }
            )
        return {
            "ok": True,
            "plans": plans,
            "providers": providers,
            "capabilities": {
                "yookassa_ready": bool(settings.yookassa_shop_id and settings.yookassa_secret_key),
                "yookassa_recurring_available": bool(settings.yookassa_enable_saved_payment_method),
                "autopay_available": bool(settings.yookassa_shop_id and settings.yookassa_secret_key and settings.yookassa_enable_saved_payment_method),
            },
        }

    @app.get("/api/economics/plans")
    async def economics_plans(x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        plans = plan_catalog(settings)
        return {
            "ok": True,
            "profit_guard_enabled": settings.profit_guard_enabled,
            "plans": [plan_economics(settings, plan) for plan in plans.values()],
            "notes": [
                "gross -> net subtracts payment fee, tax and refund-risk reserve",
                "max_ai_budget is the hard monthly AI budget behind each plan",
                "credit charge is increased automatically when the model/token cost grows",
            ],
        }

    @app.get("/api/economics/credit-packs")
    async def economics_credit_packs(x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        packs = features.credit_packs()
        return {"ok": True, "packs": [credit_pack_margin(settings, pack) for pack in packs]}

    @app.get("/api/billing/status")
    async def billing_status(telegram_user_id: int | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        telegram_id = resolve_telegram_id(user, telegram_user_id)
        status = await db.billing_status(
            telegram_id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        return {"ok": True, **status}

    @app.get("/api/billing/payout-method")
    async def get_payout_method(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        method = await db.get_payout_method(user.id)
        return {"ok": True, "payout_method": method, "safe_storage": True}

    @app.post("/api/billing/payout-method")
    async def save_payout_method(payload: PayoutMethodRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        bik, account = validate_payout_payload(payload)
        # Храним полный номер счёта только в зашифрованном виде. Во frontend отдаём только маску.
        saved = await db.upsert_payout_method(
            user.id,
            bik=bik,
            account_number_encrypted=encrypt_text(account, settings.app_secret) or "",
            account_last4=account[-4:],
            account_mask=mask_account(account),
            bank_name=(payload.bank_name or "").strip() or None,
            holder_name=(payload.holder_name or "").strip() or None,
        )
        return {"ok": True, "payout_method": saved, "safe_storage": True}

    @app.delete("/api/billing/payout-method")
    async def delete_payout_method(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await db.delete_payout_method(user.id)}

    @app.get("/api/billing/autopay")
    async def get_autopay(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {
            "ok": True,
            "autopay": await db.get_autopay_settings(user.id),
            "available": bool(settings.yookassa_shop_id and settings.yookassa_secret_key and settings.yookassa_enable_saved_payment_method),
            "reason": None if settings.yookassa_enable_saved_payment_method else "recurring_disabled",
        }

    @app.post("/api/billing/autopay")
    async def update_autopay(payload: AutopaySettingsRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        if payload.provider != "yookassa":
            return {"ok": False, "error": "Автоподписка сейчас поддерживается только через ЮKassa."}
        if payload.enabled and not settings.yookassa_enable_saved_payment_method:
            return {
                "ok": False,
                "error": "Автопродление сейчас выключено: у магазина ЮKassa не подключены recurring-платежи. Разовая покупка подписки работает без этого.",
                "requires_recurring_permission": True,
            }
        plan_key = (payload.plan or "").lower() or None
        if payload.enabled:
            plans = plan_catalog(settings)
            if not plan_key or plan_key not in plans or plan_key == "free":
                return {"ok": False, "error": "Выберите платный тариф для автоподписки."}
            existing_secret = await db.get_autopay_settings_secret(user.id)
            if not existing_secret or not existing_secret.get("payment_method_id_encrypted"):
                autopay = await db.upsert_autopay_settings(user.id, enabled=True, provider="yookassa", plan=plan_key, status="pending_payment_method")
                return {
                    "ok": False,
                    "requires_payment_method": True,
                    "error": "Сначала оплатите тариф картой/СБП с включённой автоподпиской — FounderPilot сохранит только безопасный токен ЮKassa, не данные карты.",
                    "autopay": autopay,
                }
        autopay = await db.upsert_autopay_settings(
            user.id,
            enabled=payload.enabled,
            provider="yookassa",
            plan=plan_key,
            status="active" if payload.enabled else "disabled",
        )
        return {"ok": True, "autopay": autopay}

    @app.delete("/api/billing/autopay/payment-method")
    @app.post("/api/billing/autopay/unlink")
    async def unlink_autopay_payment_method(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        autopay = await db.unlink_autopay_payment_method(user.id)
        await notify_user(user.id, title="Автопродление отключено", body="Сохранённая карта отвязана от FounderPilot. Разовые покупки подписки остаются доступными.", type="billing", action_url="subscription")
        return {"ok": True, "autopay": autopay}

    @app.post("/api/billing/autopay/run-due")
    async def run_due_autopay(limit: int = 25, x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        if not settings.yookassa_enable_saved_payment_method:
            return {"ok": False, "error": "Автосписания ЮKassa выключены в настройках сервера.", "results": []}
        results: list[dict[str, Any]] = []
        plans = plan_catalog(settings)
        for item in await db.list_due_autopay(limit=limit):
            telegram_id = int(item.get("telegram_id") or item.get("telegram_user_id"))
            plan_key = str(item.get("plan") or "").lower()
            plan = plans.get(plan_key)
            encrypted_method = item.get("payment_method_id_encrypted")
            payment_method_id = decrypt_text(encrypted_method, settings.app_secret)
            if not plan or plan.key == "free" or not payment_method_id:
                await db.mark_autopay_charge(telegram_id, status="disabled")
                results.append({"telegram_user_id": telegram_id, "ok": False, "error": "invalid_autopay_config"})
                continue
            order_id = make_order_id()
            amount, currency = price_for_provider(plan, "yookassa")
            order = await db.create_billing_order(
                order_id,
                telegram_id,
                plan.key,
                "yookassa",
                float(amount),
                currency,
                metadata={"plan_title": plan.title, "auto_renewal_charge": True},
            )
            try:
                external_id, status, raw_payment = await create_yookassa_autopayment(
                    settings,
                    payment_method_id=payment_method_id,
                    order=order,
                    plan=plan,
                )
                await db.update_billing_order(order_id, external_payment_id=external_id, payload=json.dumps(raw_payment, ensure_ascii=False), status=status)
                if status == "succeeded":
                    await activate_subscription(db, telegram_id, plan.key, "yookassa", order_id, plan.daily_limit, plan.monthly_limit)
                    await notify_user(telegram_id, title="Подписка продлена", body=f"Тариф {plan.title} продлён автоматически.", type="billing", action_url="subscription", metadata={"order_id": order_id, "plan": plan.key})
                    await db.mark_autopay_charge(telegram_id, status="active", next_charge_at=subscription_next_charge_at())
                else:
                    await db.mark_autopay_charge(telegram_id, status="retry")
                results.append({"telegram_user_id": telegram_id, "order_id": order_id, "status": status})
            except Exception as exc:  # noqa: BLE001
                await db.mark_autopay_charge(telegram_id, status="retry")
                await db.log_error("autopay", str(exc), telegram_id)
                results.append({"telegram_user_id": telegram_id, "ok": False, "error": str(exc)})
        return {"ok": True, "results": results}

    @app.post("/api/billing/create-order")
    @app.post("/api/billing/checkout")
    @app.post("/api/billing/orders")
    @app.post("/api/subscription/checkout")
    async def billing_create_order(request: Request, payload: BillingOrderRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        telegram_id = resolve_telegram_id(user, payload.telegram_user_id)
        await db.upsert_user(telegram_id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        plans = plan_catalog(settings)
        plan = plans.get(normalize_plan_key(payload.plan))
        if not plan or plan.key == "free":
            return {"ok": False, "error": "Выберите платный тариф."}
        try:
            provider = resolve_payment_provider(
                settings,
                payload.provider,
                telegram_webapp=bool(request.headers.get("x-telegram-init-data")),
            )
            amount, currency = price_for_provider(plan, provider)
        except BillingError as exc:
            return {"ok": False, "error": str(exc)}
        if amount <= 0:
            return {"ok": False, "error": "Для выбранного способа оплаты не указана цена тарифа."}

        effective_auto_renew = bool(payload.auto_renew and provider == "yookassa" and settings.yookassa_enable_saved_payment_method)
        order_id = make_order_id()
        order = await db.create_billing_order(
            order_id,
            telegram_id,
            plan.key,
            provider,
            float(amount),
            currency,
            metadata={"plan_title": plan.title, "auto_renew": effective_auto_renew, "requested_auto_renew": bool(payload.auto_renew), "requested_provider": payload.provider or "auto"},
        )

        try:
            if provider == "telegram_stars":
                bot = Bot(token=settings.bot_token)
                invoice_link = await create_telegram_stars_invoice(settings, bot, order, plan)
                await bot.session.close()
                await db.update_billing_order(order_id, payment_url=invoice_link, payload=order_id)
                return {"ok": True, "order_id": order_id, "provider": provider, "payment_url": invoice_link}
            if provider == "yookassa":
                external_id, payment_url = await create_yookassa_payment(settings, order, plan, save_payment_method=effective_auto_renew)
                await db.update_billing_order(order_id, external_payment_id=external_id, payment_url=payment_url)
                await db.record_payment(
                    order_id=order_id,
                    telegram_id=telegram_id,
                    provider="yookassa",
                    plan=plan.key,
                    amount=float(amount),
                    currency=currency,
                    status="pending",
                    external_payment_id=external_id,
                )
                return {"ok": True, "order_id": order_id, "provider": provider, "payment_url": payment_url, "auto_renew_enabled": effective_auto_renew}
            if provider == "ton":
                ton_transaction = build_ton_transaction(settings, order)
                ton_payment_url = build_ton_payment_link(settings, order)
                await db.update_billing_order(order_id, payload=order_id, metadata={"ton_transaction": ton_transaction})
                return {
                    "ok": True,
                    "order_id": order_id,
                    "provider": provider,
                    "ton_transaction": ton_transaction,
                    "payment_url": ton_payment_url,
                }
            if provider == "btcpay_btc":
                external_id, payment_url = await create_btcpay_invoice(settings, order, plan)
                await db.update_billing_order(order_id, external_payment_id=external_id, payment_url=payment_url)
                await db.record_payment(
                    order_id=order_id,
                    telegram_id=telegram_id,
                    provider="btcpay_btc",
                    plan=plan.key,
                    amount=float(amount),
                    currency=currency,
                    status="pending",
                    external_payment_id=external_id,
                )
                return {"ok": True, "order_id": order_id, "provider": provider, "payment_url": payment_url}
        except BillingError as exc:
            await db.update_billing_order(order_id, status="failed")
            await db.log_error("billing_create_order", str(exc), telegram_id)
            return {"ok": False, "error": str(exc), "order_id": order_id}
        except Exception as exc:  # noqa: BLE001
            await db.update_billing_order(order_id, status="failed")
            await db.log_error("billing_create_order", str(exc), telegram_id)
            return {"ok": False, "error": f"Не удалось создать оплату: {exc}", "order_id": order_id}

        return {"ok": False, "error": "Способ оплаты пока не поддерживается."}

    @app.get("/api/billing/order/{order_id}")
    async def billing_order(order_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        order = await db.get_billing_order(order_id)
        if not order:
            raise HTTPException(status_code=404, detail="Заказ не найден.")
        if not settings.allow_dev_auth and str(order.get("telegram_user_id")) != str(user.id):
            raise HTTPException(status_code=403, detail="Этот заказ принадлежит другому пользователю.")

        # Reliable post-payment flow: if YooKassa already accepted the money but
        # the webhook has not reached us yet, polling this endpoint finishes activation.
        if order.get("provider") == "yookassa" and str(order.get("status") or "").lower() == "pending" and order.get("external_payment_id"):
            try:
                order = await finalize_yookassa_order(order)
            except BillingError:
                # Payment may still be waiting for user confirmation. Keep order pending.
                order = await db.get_billing_order(order_id) or order
            except Exception as exc:  # noqa: BLE001
                await db.log_error("billing_order_yookassa_verify", str(exc), user.id)
                order = await db.get_billing_order(order_id) or order
        return {"ok": True, "order": order, "status": order.get("status")}

    @app.post("/api/billing/ton/verify")
    async def ton_verify(payload: TonVerifyRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        telegram_id = resolve_telegram_id(user, payload.telegram_user_id)
        order = await db.get_billing_order(payload.order_id)
        if not order or order.get("provider") != "ton":
            return {"ok": False, "error": "TON-заказ не найден."}
        if order.get("telegram_user_id") != str(telegram_id):
            return {"ok": False, "error": "Этот заказ принадлежит другому пользователю."}
        if order.get("status") == "paid":
            return {"ok": True, "status": "paid"}
        ok, message = await best_effort_ton_verify(settings, order, payload.tx_hash)
        if not ok:
            return {"ok": False, "error": message, "status": "pending"}
        plan = plan_catalog(settings)[str(order["plan"])]
        result = await activate_subscription(db, telegram_id, plan.key, "ton", order["id"], plan.daily_limit, plan.monthly_limit)
        await db.record_payment(
            order_id=order["id"],
            telegram_id=telegram_id,
            provider="ton",
            plan=plan.key,
            amount=float(order["amount"]),
            currency="TON",
            status="paid",
            external_payment_id=payload.tx_hash,
            raw_event={"wallet_address": payload.wallet_address, "tx_hash": payload.tx_hash},
        )
        return {"ok": True, "status": "paid", "subscription": result}

    @app.post("/api/billing/webhooks/yookassa")
    @app.post("/api/billing/yookassa/webhook")
    async def yookassa_webhook(request: Request) -> dict[str, Any]:
        raw = await read_json_dict(request)
        obj = raw.get("object", {}) if isinstance(raw, dict) else {}
        metadata = obj.get("metadata", {}) or {}
        order_id = metadata.get("order_id")
        external_id = obj.get("id")
        event_type = raw.get("event")
        await db.record_payment_event("yookassa", raw, event_type, order_id, external_id)
        kind = str(metadata.get("kind") or "").lower()
        if kind == "credit_pack":
            credit_order_id = metadata.get("credit_pack_order_id") or order_id
            credit_order = await features.get_credit_pack_order(str(credit_order_id)) if credit_order_id else None
            if not credit_order and external_id:
                credit_order = await features.find_credit_pack_order_by_external_id(str(external_id))
            if not credit_order:
                return {"ok": True, "ignored": True, "kind": "credit_pack"}
            status = obj.get("status")
            if event_type == "payment.succeeded" or status == "succeeded":
                if str(credit_order.get("status") or "").lower() == "paid":
                    return {"ok": True, "already_processed": True, "kind": "credit_pack"}
                await finalize_credit_pack_order(credit_order, raw_event=raw)
            elif status == "canceled" or event_type == "payment.canceled":
                try:
                    verified = await fetch_yookassa_payment(settings, str(external_id))
                except BillingError as exc:
                    await db.log_error("yookassa_credit_pack_cancel_verify", str(exc))
                    return {"ok": True, "ignored": True, "kind": "credit_pack"}
                if verified.get("status") == "canceled":
                    await features.update_credit_pack_order_payment(credit_order["id"], status="canceled")
            return {"ok": True, "kind": "credit_pack"}

        order = await db.get_billing_order(order_id) if order_id else None
        if not order and external_id:
            order = await db.find_billing_order_by_external_id(external_id)
        if not order:
            return {"ok": True, "ignored": True}
        status = obj.get("status")
        if event_type == "payment.succeeded" or status == "succeeded":
            if str(order.get("status") or "").lower() == "paid":
                return {"ok": True, "already_processed": True}
            await finalize_yookassa_order(order, raw_event=raw)
        elif status == "canceled" or event_type == "payment.canceled":
            try:
                verified = await fetch_yookassa_payment(settings, str(external_id))
            except BillingError as exc:
                await db.log_error("yookassa_cancel_verify", str(exc), int(order["telegram_user_id"]))
                return {"ok": True, "ignored": True}
            if verified.get("status") == "canceled":
                await db.update_billing_order(order["id"], status="canceled")
        return {"ok": True}

    @app.post("/api/billing/webhooks/btcpay")
    async def btcpay_webhook(request: Request, btcpay_sig: str | None = Header(default=None, alias="BTCPay-Sig")) -> dict[str, Any]:
        raw_body = await request.body()
        if len(raw_body) > settings.max_request_body_bytes:
            raise HTTPException(status_code=413, detail="Request body is too large.")
        if not verify_btcpay_signature(settings, raw_body, btcpay_sig):
            raise HTTPException(status_code=403, detail="Неверная подпись BTCPay.")
        raw = json_loads_safe(raw_body)
        invoice_id = raw.get("invoiceId") or raw.get("id") or raw.get("data", {}).get("id")
        event_type = raw.get("type") or raw.get("event")
        order = await db.find_billing_order_by_external_id(str(invoice_id)) if invoice_id else None
        await db.record_payment_event("btcpay_btc", raw, event_type, order.get("id") if order else None, str(invoice_id) if invoice_id else None)
        if not order:
            return {"ok": True, "ignored": True}
        if str(order.get("status") or "").lower() == "paid":
            return {"ok": True, "already_processed": True}
        event_text = str(event_type or "").lower()
        status_text = str(raw.get("status") or raw.get("invoiceStatus") or "").lower()
        if any(word in event_text or word in status_text for word in ["settled", "confirmed", "complete", "paid"]):
            plan = plan_catalog(settings)[str(order["plan"])]
            await activate_subscription(db, order["telegram_user_id"], plan.key, "btcpay_btc", order["id"], plan.daily_limit, plan.monthly_limit)
            await notify_user(order["telegram_user_id"], title="Подписка активна", body=f"Тариф {plan.title} успешно подключён.", type="billing", action_url="subscription", metadata={"order_id": order["id"], "plan": plan.key})
            await db.record_payment(
                order_id=order["id"],
                telegram_id=int(order["telegram_user_id"]),
                provider="btcpay_btc",
                plan=plan.key,
                amount=float(order["amount"]),
                currency=order["currency"],
                status="paid",
                external_payment_id=str(invoice_id),
                raw_event=raw,
            )
        elif any(word in event_text or word in status_text for word in ["expired", "invalid", "failed"]):
            await db.update_billing_order(order["id"], status="expired")
        return {"ok": True}

    @app.get("/api/referral")
    async def referral(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        bot_username = None
        return {"ok": True, **await db.referral_stats(user.id, bot_username)}

    @app.get("/api/stats", response_model=StatsResponse)
    async def stats(user: TelegramUser = Depends(current_user)) -> StatsResponse:
        access = await db.get_access_state(
            user.id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        return StatsResponse(
            used_today=access["used_today"],
            daily_limit=access["current_limit"],
            per_minute_limit=settings.per_minute_limit,
            free_limit=access["free_limit"],
            monthly_limit=access["monthly_limit"],
            used_total=access["used_total"],
            used_period=access["used_period"],
            remaining=access["remaining"],
            plan=access["plan"],
            status=access["status"],
            status_label=access["status_label"],
            subscription_until=access["subscription_until"],
            unlimited=access["unlimited"],
        )

    @app.get("/api/profile")
    async def profile(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        access = await db.get_access_state(
            user.id,
            free_limit_default=settings.free_trial_requests,
            monthly_limit_default=settings.free_monthly_credits,
        )
        business = await db.get_business_profile(user.id)
        return {
            "user": {
                "id": user.id,
                "username": user.username,
                "first_name": user.first_name,
                "last_name": user.last_name,
                "business_profile": (business or {}).get("description") or (business or {}).get("main_problem") or "",
                "company_name": (business or {}).get("business_name") or (business or {}).get("niche") or "",
            },
            "business_profile": business,
            "usage": usage_payload(access),
            "limits": {
                "free_trial_requests": settings.free_trial_requests,
                "subscriber_monthly_limit": settings.subscriber_monthly_limit,
                "per_minute_limit": settings.per_minute_limit,
            },
        }

    @app.post("/api/profile")
    async def update_profile(request: Request, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        payload = await read_json_dict(request)
        company = str(payload.get("company_name") or payload.get("business_name") or "").strip()
        description = str(payload.get("business_profile") or payload.get("description") or "").strip()
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        saved = await db.upsert_business_profile(user.id, {
            "business_name": company,
            "description": description,
            "main_problem": description,
        })
        return {"ok": True, "profile": saved}


    @app.get("/api/projects")
    async def projects(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_projects(user.id), "active": await features.get_active_project(user.id)}

    @app.post("/api/projects")
    async def create_project(payload: ProjectRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        project = await features.create_project(user.id, payload.model_dump(exclude_none=True))
        return {"ok": True, "project": project}

    @app.get("/api/projects/current")
    async def current_project(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "project": await features.get_active_project(user.id)}

    @app.get("/api/projects/{project_id}")
    async def get_project(project_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        project = await features.get_project(user.id, project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Проект не найден.")
        return {"ok": True, "project": project, "memory": await features.list_memory(user.id, project_id)}

    @app.patch("/api/projects/{project_id}")
    async def update_project(project_id: str, payload: ProjectRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        project = await features.update_project(user.id, project_id, payload.model_dump(exclude_none=True))
        if not project:
            raise HTTPException(status_code=404, detail="Проект не найден.")
        return {"ok": True, "project": project}

    @app.post("/api/projects/{project_id}/activate")
    async def activate_project(project_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        project = await features.update_project(user.id, project_id, {"is_active": True})
        if not project:
            raise HTTPException(status_code=404, detail="Проект не найден.")
        return {"ok": True, "project": project}

    @app.delete("/api/projects/{project_id}")
    async def delete_project(project_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await features.delete_project(user.id, project_id)}

    @app.get("/api/memory")
    async def memory(project_id: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_memory(user.id, project_id)}

    @app.post("/api/memory")
    async def add_memory(payload: MemoryRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        data = payload.model_dump(exclude_none=True)
        if not (data.get("value") or data.get("content")):
            return {"ok": False, "error": "Память не может быть пустой."}
        return {"ok": True, "item": await features.add_memory(user.id, data)}

    @app.delete("/api/memory/{memory_id}")
    async def delete_memory(memory_id: int, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await features.delete_memory(user.id, memory_id)}

    @app.get("/api/templates")
    async def templates(category: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_templates(user.id, category)}

    @app.post("/api/templates")
    async def create_template(payload: TemplateRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "template": await features.create_template(user.id, payload.model_dump())}

    @app.delete("/api/templates/{template_id}")
    async def delete_template(template_id: int, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await features.delete_template(user.id, template_id)}


    @app.get("/api/startup-suite")
    async def startup_suite(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        active = await features.get_active_project(user.id)
        project_id = active.get("id") if active else None
        return {
            "ok": True,
            "active_project": active,
            "projects": await features.list_projects(user.id),
            "memory": await features.list_memory(user.id, project_id),
            "documents": await features.list_documents(user.id, project_id),
            "roadmaps": await features.list_roadmaps(user.id, project_id),
            "tasks": await features.list_tasks(user.id, project_id),
            "score": await features.latest_project_score(user.id, project_id),
            "finance": await features.project_finance_snapshot(user.id, project_id),
            "analytics": await features.analytics_summary(user.id),
        }

    @app.get("/api/projects/{project_id}/finance")
    async def project_finance(project_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "finance": await features.project_finance_snapshot(user.id, project_id)}

    @app.get("/api/documents")
    async def documents(project_id: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_documents(user.id, project_id)}

    @app.post("/api/documents")
    async def create_document(payload: DocumentRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "document": await features.create_document(user.id, payload.model_dump(exclude_none=True))}

    @app.get("/api/documents/{document_id}")
    async def get_document(document_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        doc = await features.get_document(user.id, document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Документ не найден.")
        return {"ok": True, "document": doc}

    @app.delete("/api/documents/{document_id}")
    async def delete_document(document_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await features.delete_document(user.id, document_id)}

    @app.post("/api/documents/generate")
    async def generate_document(payload: DocumentGenerateRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        await db.expire_subscription_if_needed(user.id)
        project = await features.get_project(user.id, payload.project_id) if payload.project_id else await features.get_active_project(user.id)
        project_context = await features.project_context_text(user.id)
        doc_type = (payload.document_type or "business_plan").strip().lower()
        mode_map = {
            "business_plan": "strategy",
            "marketing_plan": "sales_plan",
            "financial_model": "unit",
            "pitch_deck": "pitch",
            "competitor_analysis": "competitor_analysis",
            "content_plan": "content_plan",
            "launch_plan": "next_step",
            "legal_checklist": "strategy",
        }
        mode = payload.mode or mode_map.get(doc_type, "strategy")
        title_map = {
            "business_plan": "Бизнес-план",
            "marketing_plan": "Маркетинг-план",
            "financial_model": "Финансовая модель",
            "pitch_deck": "Pitch deck",
            "competitor_analysis": "Анализ конкурентов",
            "content_plan": "Контент-план",
            "launch_plan": "План запуска",
            "legal_checklist": "Юридический чек-лист",
        }
        title = payload.title or title_map.get(doc_type, "Документ")
        user_prompt = (payload.prompt or "").strip()
        if not user_prompt:
            if project_context:
                user_prompt = f"Подготовь документ типа: {title}. Используй контекст проекта и сделай результат максимально прикладным."
            else:
                user_prompt = f"Подготовь документ типа: {title}. Если данных не хватает, сделай качественный черновик и отметь допущения."
        profile = await db.get_user_profile(user.id)
        plan_key = str((profile or {}).get("plan") or "free").lower()
        ai_model = ai_client.model_for_plan(plan_key)
        access_state = await db.get_access_state(user.id, free_limit_default=settings.free_trial_requests, monthly_limit_default=settings.free_monthly_credits)
        optional_fields = {
            "document_type": doc_type,
            "project_memory": project_context,
            "project": json.dumps(project or {}, ensure_ascii=False),
            "output_policy": "Верни структурированный документ в Markdown с оглавлением, таблицами где полезно, конкретными шагами и допущениями.",
        }
        estimate = estimate_credits(mode, user_prompt, model=ai_model, optional_fields=optional_fields)
        estimate = await apply_profit_guard_to_estimate(user.id, estimate)
        try:
            await rate_limiter.check(user.id, estimate.credits)
            await db.reserve_credits(
                user.id,
                estimate.request_id,
                f"document_{doc_type}",
                estimate.credits,
                free_limit_default=settings.free_trial_requests,
                monthly_limit_default=settings.free_monthly_credits,
                metadata={"endpoint": "/api/documents/generate", "model": ai_model, "document_type": doc_type},
            )
            content = await ai_client.ask_business_ai(mode, user_prompt, optional_fields, plan_key=plan_key, access_context=access_state)
            await db.finalize_credit_charge(
                user.id,
                estimate.request_id,
                f"document_{doc_type}",
                estimate.credits,
                estimate.credits,
                model=ai_model,
                input_tokens=estimate.input_tokens_estimated,
                output_tokens=estimate_output_tokens(content),
            )
        except (RateLimitError, CreditLimitError) as exc:
            raise HTTPException(status_code=429, detail=str(exc)) from exc
        except AIClientError as exc:
            await db.refund_reserved_credits(user.id, estimate.request_id, str(exc), tool_id=f"document_{doc_type}", estimated_credits=estimate.credits, model=ai_model)
            raise HTTPException(status_code=502, detail=str(exc)) from exc
        except Exception as exc:  # noqa: BLE001
            await db.refund_reserved_credits(user.id, estimate.request_id, str(exc), tool_id=f"document_{doc_type}", estimated_credits=estimate.credits, model=ai_model)
            raise
        document = await features.create_document(user.id, {
            "project_id": (project or {}).get("id") or payload.project_id,
            "title": title,
            "document_type": doc_type,
            "content": content,
            "source": "ai",
            "metadata": {"mode": mode, "model": ai_model},
        })
        return {"ok": True, "document": document, "usage": {"credits_charged": estimate.credits, "model": ai_model}}

    def _filename_safe(value: str) -> str:
        clean = re.sub(r"[^A-Za-z0-9._ -]+", "_", value, flags=re.IGNORECASE).strip(" ._")
        return (clean or "founderpilot-document")[:80]

    @app.get("/api/documents/{document_id}/export")
    async def export_document(document_id: str, format: str = "md", user: TelegramUser = Depends(current_user)) -> Response:
        doc = await features.get_document(user.id, document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Документ не найден.")
        title = str(doc.get("title") or "Документ")
        content = str(doc.get("content") or "")
        fmt = (format or "md").lower().strip()
        basename = _filename_safe(title)
        if fmt in {"md", "markdown", "txt"}:
            body = f"# {title}\n\n{content}\n"
            return Response(body, media_type="text/markdown; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{basename}.md"'})
        if fmt == "html":
            escaped = html_escape(content).replace("\n", "<br>")
            html = f"<!doctype html><html><head><meta charset='utf-8'><title>{html_escape(title)}</title><style>body{{font-family:-apple-system,BlinkMacSystemFont,Segoe UI,Arial,sans-serif;max-width:860px;margin:40px auto;line-height:1.55;padding:0 20px}}h1{{font-size:30px}}</style></head><body><h1>{html_escape(title)}</h1><div>{escaped}</div></body></html>"
            return Response(html, media_type="text/html; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{basename}.html"'})
        if fmt == "docx":
            try:
                from docx import Document  # type: ignore
                document = Document()
                document.add_heading(title, level=1)
                for block in content.split("\n"):
                    text = block.strip()
                    if not text:
                        continue
                    if text.startswith("# "):
                        document.add_heading(text[2:].strip(), level=1)
                    elif text.startswith("## "):
                        document.add_heading(text[3:].strip(), level=2)
                    elif text.startswith("- "):
                        document.add_paragraph(text[2:].strip(), style="List Bullet")
                    else:
                        document.add_paragraph(text)
                buf = BytesIO()
                document.save(buf)
                return Response(buf.getvalue(), media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f'attachment; filename="{basename}.docx"'})
            except Exception as exc:  # noqa: BLE001
                await db.log_error("document_docx_export", str(exc), user.id)
                raise HTTPException(status_code=500, detail="DOCX-экспорт временно недоступен.") from exc
        raise HTTPException(status_code=400, detail="Поддерживаются форматы md, html и docx.")

    @app.post("/api/projects/{project_id}/score")
    async def score_project(project_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        project = await features.get_project(user.id, project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Проект не найден.")
        return {"ok": True, "score": await features.score_project(user.id, project_id)}

    @app.get("/api/projects/{project_id}/score")
    async def latest_project_score(project_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "score": await features.latest_project_score(user.id, project_id)}

    @app.get("/api/roadmaps")
    async def roadmaps(project_id: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_roadmaps(user.id, project_id)}

    @app.post("/api/roadmaps")
    async def create_roadmap(payload: RoadmapRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "roadmap": await features.create_roadmap(user.id, payload.model_dump(exclude_none=True))}

    @app.get("/api/roadmaps/{roadmap_id}")
    async def get_roadmap(roadmap_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        roadmap = await features.get_roadmap(user.id, roadmap_id)
        if not roadmap:
            raise HTTPException(status_code=404, detail="Roadmap не найден.")
        return {"ok": True, "roadmap": roadmap}

    @app.get("/api/tasks")
    async def tasks(project_id: str | None = None, status: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_tasks(user.id, project_id, status)}

    @app.post("/api/tasks")
    async def upsert_task(payload: TaskRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "task": await features.upsert_task(user.id, payload.model_dump(exclude_none=True))}

    @app.delete("/api/tasks/{task_id}")
    async def delete_task(task_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": await features.delete_task(user.id, task_id)}

    @app.post("/api/ai/feedback")
    async def ai_feedback(payload: AiFeedbackRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        item = await features.save_ai_feedback(user.id, payload.model_dump(exclude_none=True))
        return {"ok": True, "feedback": item}

    @app.get("/api/credits/packs")
    async def credit_packs() -> dict[str, Any]:
        packs = []
        for pack in features.credit_packs():
            packs.append({**pack, "profit_guard": credit_pack_margin(settings, pack)})
        return {"ok": True, "packs": packs}

    async def finalize_credit_pack_order(order: dict[str, Any], raw_event: dict[str, Any] | None = None) -> dict[str, Any]:
        if not order:
            raise BillingError("Заказ пакета кредитов не найден.")
        if str(order.get("status") or "").lower() == "paid":
            return {"ok": True, "status": "paid", "credits": int(order.get("credits") or 0), "order": order}
        external_id = str(order.get("external_payment_id") or "")
        if not external_id:
            raise BillingError("У заказа пакета нет ID платежа ЮKassa.")
        # Verify amount/currency/kind manually, then grant credits exactly once.
        payment = await verify_yookassa_payment_for_credit_pack(order, external_id)
        result = await features.grant_credit_pack(int(order["telegram_user_id"]), order["id"], "yookassa")
        await features.update_credit_pack_order_payment(order["id"], status="paid", provider="yookassa", external_payment_id=external_id)
        await db.record_payment(
            order_id=order["id"],
            telegram_id=int(order["telegram_user_id"]),
            provider="yookassa_credit_pack",
            plan=str(order.get("pack_key") or "credit_pack"),
            amount=float(order.get("amount") or 0),
            currency=order.get("currency") or "RUB",
            status="succeeded",
            external_payment_id=external_id,
            raw_event=raw_event or payment,
        )
        return {"ok": True, "status": "paid", "credits": result.get("credits"), "order": await features.get_credit_pack_order(order["id"])}

    async def verify_yookassa_payment_for_credit_pack(order: dict[str, Any], payment_id: str) -> dict[str, Any]:
        from app.billing import fetch_yookassa_payment, rub_value
        from decimal import Decimal
        data = await fetch_yookassa_payment(settings, payment_id)
        metadata = data.get("metadata") or {}
        amount = data.get("amount") or {}
        if str(metadata.get("kind") or "") != "credit_pack":
            raise BillingError("ЮKassa: платёж не относится к пакету кредитов.")
        if str(metadata.get("credit_pack_order_id") or metadata.get("order_id") or "") != str(order["id"]):
            raise BillingError("ЮKassa: metadata.order_id не совпадает с заказом пакета.")
        if str(amount.get("currency") or "").upper() != "RUB":
            raise BillingError("ЮKassa: валюта пакета кредитов не совпадает.")
        if str(amount.get("value") or "") != rub_value(Decimal(str(order["amount"]))):
            raise BillingError("ЮKassa: сумма пакета кредитов не совпадает.")
        if data.get("status") != "succeeded":
            raise BillingError("ЮKassa: платёж пакета кредитов ещё не подтверждён.")
        return data

    @app.post("/api/credits/packs/order")
    async def create_credit_pack_order(payload: CreditPackOrderRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        telegram_id = resolve_telegram_id(user, payload.telegram_user_id)
        await db.upsert_user(telegram_id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        pack_key = payload.pack_key or payload.pack_id
        if not pack_key:
            return {"ok": False, "error": "Пакет кредитов не выбран."}
        provider = (payload.provider or "yookassa").strip().lower()
        pack = next((p for p in features.credit_packs() if p["key"] == pack_key or p["id"] == pack_key), None)
        if not pack:
            return {"ok": False, "error": "Пакет кредитов не найден."}
        margin = credit_pack_margin(settings, pack, provider="yookassa")
        if not margin.get("is_profitable"):
            return {"ok": False, "error": "Пакет временно отключён: цена не покрывает безопасный лимит AI-расходов.", "economics": margin}
        order = await features.create_credit_pack_order(telegram_id, pack_key, provider)
        if provider != "yookassa":
            return {**order, "disabled": True, "error": "Пакеты кредитов сейчас оплачиваются через ЮKassa."}
        if not provider_enabled(settings, "yookassa"):
            return {"ok": False, "error": "ЮKassa пока не настроена."}
        try:
            payment_order = {
                "id": order["order_id"],
                "telegram_user_id": str(telegram_id),
                "pack_key": pack_key,
                "amount": float((pack or {}).get("amount") or 0),
                "currency": "RUB",
            }
            external_id, payment_url = await create_yookassa_credit_pack_payment(settings, payment_order, pack or {})
            updated = await features.update_credit_pack_order_payment(order["order_id"], status="pending", provider="yookassa", external_payment_id=external_id, payment_url=payment_url)
            return {"ok": True, "order_id": order["order_id"], "provider": "yookassa", "payment_url": payment_url, "order": updated, "pack": pack}
        except BillingError as exc:
            await features.update_credit_pack_order_payment(order["order_id"], status="failed", provider="yookassa")
            return {"ok": False, "error": str(exc), "order_id": order["order_id"]}
        except Exception as exc:  # noqa: BLE001
            await features.update_credit_pack_order_payment(order["order_id"], status="failed", provider="yookassa")
            await db.log_error("credit_pack_yookassa_order", str(exc), telegram_id)
            return {"ok": False, "error": "Не удалось создать оплату пакета кредитов. Попробуйте позже.", "order_id": order["order_id"]}

    @app.get("/api/credits/packs/order/{order_id}")
    async def get_credit_pack_order(order_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        order = await features.get_credit_pack_order(order_id)
        if not order:
            raise HTTPException(status_code=404, detail="Заказ пакета кредитов не найден.")
        if not settings.allow_dev_auth and str(order.get("telegram_user_id")) != str(user.id):
            raise HTTPException(status_code=403, detail="Этот заказ принадлежит другому пользователю.")
        if order.get("provider") == "yookassa" and str(order.get("status") or "").lower() == "pending" and order.get("external_payment_id"):
            try:
                await finalize_credit_pack_order(order)
                order = await features.get_credit_pack_order(order_id) or order
            except BillingError:
                pass
        return {"ok": True, "order": order, "status": order.get("status")}

    @app.post("/api/credits/packs/{order_id}/grant")
    async def grant_credit_pack(order_id: str, telegram_user_id: int, x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        return await features.grant_credit_pack(telegram_user_id, order_id, "admin_grant")

    @app.get("/api/analytics/summary")
    async def analytics_summary(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "summary": await features.analytics_summary(user.id)}

    @app.get("/api/export/history.txt")
    async def export_history_txt(user: TelegramUser = Depends(current_user)) -> PlainTextResponse:
        rows = await db.list_recent_requests(user.id, limit=50)
        lines = ["FounderPilot export", ""]
        for row in rows:
            lines.append(f"[{row.get('created_at')}] {row.get('mode')}")
            lines.append(str(row.get("user_text") or row.get("text") or ""))
            lines.append(str(row.get("ai_answer") or row.get("answer") or ""))
            lines.append("---")
        return PlainTextResponse("\n".join(lines), media_type="text/plain; charset=utf-8")

    @app.get("/api/app/meta")
    async def app_meta() -> dict[str, Any]:
        return {
            "ok": True,
            "version": settings.app_version,
            "updated_at": settings.app_updated_at,
            "changelog": settings.app_changelog,
            "support": settings.support_public_name,
        }

    @app.get("/api/notifications/preferences")
    async def notification_preferences(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "preferences": await features.notification_preferences(user.id)}

    @app.post("/api/notifications/preferences")
    async def update_notification_preferences(payload: NotificationPrefsRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "preferences": await features.update_notification_preferences(user.id, payload.model_dump(exclude_none=True))}

    @app.get("/api/notifications")
    async def notifications(unread: bool = False, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        items = await features.list_notifications(user.id, unread_only=unread)
        unread_count = len([item for item in items if not item.get("read_at")]) if not unread else len(items)
        return {"ok": True, "items": items, "unread_count": unread_count}

    @app.post("/api/notifications/read")
    async def notifications_read(request: Request, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        payload = await read_json_dict(request)
        ids = payload.get("ids")
        if ids is not None and not isinstance(ids, list):
            raise HTTPException(status_code=422, detail="ids должен быть списком.")
        changed = await features.mark_notifications_read(user.id, ids)
        return {"ok": True, "updated": changed}


    @app.get("/api/support/tickets")
    async def support_tickets(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.list_support_tickets(user.id)}

    @app.post("/api/support/tickets")
    async def create_support_ticket(payload: SupportTicketRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        await db.upsert_user(user.id, user.username, user.first_name, user.last_name, photo_url=user.photo_url)
        user_profile = await db.get_user_profile(user.id) or {}
        subject = (payload.subject or "").strip()
        if not subject:
            compact = " ".join(payload.message.strip().split())
            subject = compact[:80] or "Обращение в поддержку"
        ticket = await features.create_support_ticket(
            user.id,
            subject=subject,
            message=payload.message,
            category=payload.category,
            user_name=" ".join(part for part in [user.first_name, user.last_name] if part).strip() or user.username or str(user.id),
            username=user.username,
            plan=str(user_profile.get("plan") or "free"),
        )
        group_sent = await send_support_ticket_to_group(ticket, payload.message)
        messages = await features.list_support_messages(ticket["id"])
        await notify_user(user.id, title="Обращение создано", body=f"Тикет {ticket['id']} отправлен в поддержку.", type="support", action_url="profile:support", metadata={"ticket_id": ticket["id"]})
        return {"ok": True, "ticket": ticket, "messages": messages, "group_sent": group_sent}

    @app.get("/api/support/tickets/{ticket_id}")
    async def support_ticket(ticket_id: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        ticket = await features.get_support_ticket_for_user(user.id, ticket_id)
        if not ticket:
            raise HTTPException(status_code=404, detail="Тикет поддержки не найден.")
        return {"ok": True, "ticket": ticket, "messages": await features.list_support_messages(ticket_id)}

    @app.post("/api/support/tickets/{ticket_id}/messages")
    async def add_support_ticket_message(ticket_id: str, payload: SupportMessageRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        ticket = await features.get_support_ticket_for_user(user.id, ticket_id)
        if not ticket:
            raise HTTPException(status_code=404, detail="Тикет поддержки не найден.")
        if str(ticket.get("status") or "") == "closed":
            ticket = await features.set_support_ticket_status(ticket_id, "open") or ticket
        message = await features.add_support_message(
            ticket_id,
            author_type="user",
            author_telegram_id=user.id,
            author_name=user.username or user.first_name or str(user.id),
            content=payload.message,
            source="mini_app",
            status="waiting_support",
        )
        sent = await send_support_followup_to_group(ticket, payload.message)
        ticket = await features.get_support_ticket_for_user(user.id, ticket_id) or ticket
        await notify_user(user.id, title="Сообщение отправлено", body="Мы передали уточнение в поддержку.", type="support", action_url="profile:support", metadata={"ticket_id": ticket_id})
        return {"ok": True, "ticket": ticket, "message": message, "group_sent": sent}

    @app.post("/api/support/tickets/{ticket_id}/status")
    async def update_support_ticket_status(ticket_id: str, payload: SupportTicketStatusRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        ticket = await features.get_support_ticket_for_user(user.id, ticket_id)
        if not ticket:
            raise HTTPException(status_code=404, detail="Тикет поддержки не найден.")
        clean_status = payload.status if payload.status in {"open", "waiting_support", "answered", "closed"} else "closed"
        updated = await features.set_support_ticket_status(ticket_id, clean_status)
        return {"ok": True, "ticket": updated}

    @app.get("/api/organizations")
    async def organizations(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, **await features.list_organizations(user.id, user.username)}

    @app.get("/api/organizations/current")
    async def current_organization(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        data = await features.list_organizations(user.id, user.username)
        return {"ok": True, "organization": data.get("active"), **data}

    @app.get("/api/organizations/members")
    async def organization_members(organization_id: str | None = None, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        try:
            return {"ok": True, "items": await features.list_organization_members(user.id, organization_id)}
        except PermissionError as exc:
            return {"ok": False, "error": str(exc)}

    @app.post("/api/organizations")
    async def create_organization(payload: OrganizationRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        try:
            org = await features.create_organization(user.id, payload.model_dump(exclude_none=True))
            return {"ok": True, "organization": org}
        except PermissionError as exc:
            return {"ok": False, "error": str(exc)}

    @app.post("/api/organizations/{organization_id}/invites")
    async def invite_organization_member(organization_id: str, payload: OrganizationInviteRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        try:
            invite = await features.invite_organization_member(user.id, organization_id, payload.username)
            sent = await send_organization_invite(invite)
            return {"ok": True, "invite": invite, "notification_sent": sent}
        except (PermissionError, ValueError) as exc:
            return {"ok": False, "error": str(exc)}

    @app.post("/api/organizations/invite")
    async def invite_organization_member_fallback(payload: OrganizationInviteRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        data = await features.list_organizations(user.id, user.username)
        owned = data.get("owned") or []
        if not owned:
            return {"ok": False, "error": "Сначала создайте организацию."}
        return await invite_organization_member(str(owned[0]["id"]), payload, user)

    @app.get("/api/organizations/invites/pending")
    async def pending_organization_invites(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        return {"ok": True, "items": await features.pending_invites(user.id, user.username)}

    @app.post("/api/organizations/invites/{token}/accept")
    async def accept_organization_invite(token: str, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        try:
            result = await features.accept_organization_invite(user.id, user.username, token)
            return result
        except (PermissionError, ValueError) as exc:
            return {"ok": False, "error": str(exc)}

    @app.post("/api/organizations/invites/accept")
    async def accept_organization_invite_body(payload: OrganizationAcceptRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        try:
            result = await features.accept_organization_invite(user.id, user.username, payload.token)
            return result
        except (PermissionError, ValueError) as exc:
            return {"ok": False, "error": str(exc)}

    @app.post("/api/organizations/invites/decline")
    async def decline_organization_invite_body(payload: OrganizationAcceptRequest, user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        try:
            return await features.decline_organization_invite(user.id, user.username, payload.token)
        except (PermissionError, ValueError) as exc:
            return {"ok": False, "error": str(exc)}

    @app.get("/api/admin/overview")
    async def admin_overview(x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        return {"ok": True, **await features.admin_overview()}

    @app.get("/api/admin/users")
    @app.get("/api/admin/owner-overview")
    async def admin_owner_overview(user: TelegramUser = Depends(current_user)) -> dict[str, Any]:
        if not settings.is_admin(user.id):
            raise HTTPException(status_code=403, detail="Недостаточно прав.")
        return {"ok": True, **await features.admin_overview()}

    async def admin_users(limit: int = 50, x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        return {"ok": True, "items": await features.admin_users(limit)}

    @app.get("/api/admin/stats")
    async def admin_stats(x_admin_secret: str | None = Header(default=None)) -> dict[str, Any]:
        require_admin(x_admin_secret)
        return {"ok": True, **await db.admin_stats()}

    return app


async def _run_bot_polling(settings: Settings) -> None:
    """Run Telegram long polling in the current process.

    In combined Railway mode this runs next to Uvicorn as a background task,
    so Mini App/API and the Telegram bot are alive at the same time and share
    the same PostgreSQL database.
    """
    db = Database(settings.database_dsn)
    await db.init()
    await init_features(settings.database_dsn)
    ai_client = OpenRouterClient(settings)
    rate_limiter = RateLimiter(settings, db)
    bot = Bot(token=settings.bot_token)
    dp = build_dispatcher(settings, db, ai_client, rate_limiter)
    try:
        logger.info("Telegram bot polling started in %s mode", settings.normalized_bot_service_mode)
        await dp.start_polling(bot)
    finally:
        await bot.session.close()


async def _run_bot_supervisor(settings: Settings) -> None:
    """Keep the bot alive without killing the web service on Telegram hiccups."""
    delay = max(1.0, float(settings.bot_restart_delay_seconds))
    while True:
        try:
            await _run_bot_polling(settings)
            logger.warning("Telegram bot polling finished unexpectedly; restarting in %.1fs", delay)
        except asyncio.CancelledError:
            raise
        except Exception as exc:  # noqa: BLE001
            logger.exception("Telegram bot polling error: %s", exc)
            if settings.bot_polling_strict and settings.normalized_bot_service_mode == "bot":
                raise
            logger.warning("Mini App/API keep running. Bot polling will retry in %.1fs", delay)
        await asyncio.sleep(delay)


async def async_main() -> None:
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(name)s: %(message)s")
    settings = get_settings()
    issues = runtime_setting_issues(settings)
    if settings.strict_runtime_validation:
        require_runtime_settings(settings)
    else:
        for issue in issues:
            logger.warning("Runtime settings warning: %s", issue)

    logger.info("FounderPilot runtime mode: %s", settings.normalized_bot_service_mode)

    if settings.normalized_bot_service_mode == "bot":
        if not settings.should_run_bot:
            raise RuntimeError("BOT_SERVICE_MODE=bot requires BOT_TOKEN")
        logger.info("FounderPilot bot-only service is running")
        await _run_bot_supervisor(settings)
        return

    app = create_app(settings)
    config = uvicorn.Config(app=app, host=settings.bind_host, port=settings.port, log_level="info")
    server = uvicorn.Server(config=config)

    logger.info("FounderPilot AI web service is running")
    logger.info("Mini App local URL: http://%s:%s/app", settings.bind_host, settings.port)
    logger.info("Telegram WebApp URL from .env: %s", settings.webapp_url)

    bot_task: asyncio.Task | None = None
    if settings.should_run_bot:
        logger.info("Telegram bot polling is enabled in the same process as Mini App/API")
        bot_task = asyncio.create_task(_run_bot_supervisor(settings), name="telegram-bot-supervisor")
    else:
        if not settings.bot_token.strip():
            logger.warning("Telegram bot polling is disabled because BOT_TOKEN is empty")
        else:
            logger.info("Telegram bot polling is disabled by BOT_SERVICE_MODE=%s", settings.normalized_bot_service_mode)

    try:
        await server.serve()
    finally:
        if bot_task is not None:
            bot_task.cancel()
            with suppress(asyncio.CancelledError):
                await bot_task


def main() -> None:
    asyncio.run(async_main())
